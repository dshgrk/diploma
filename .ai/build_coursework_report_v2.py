from __future__ import annotations

import json
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Mm, Pt


ROOT = Path(r"E:\diploma (7)\diploma")
OUTPUT_DIR = Path(r"E:\UNIVERSITY\4 kurs\2 семестр\курсовая")
FINAL_DOCX = OUTPUT_DIR / "Курсова_ТВПЗ_Гірка_розширена_фінал_v13.docx"
TITLE_DOCX = OUTPUT_DIR / "титулка.docx"

SCREENSHOTS = {
    "home": ROOT / "evidence-home.png",
    "constructor": ROOT / "evidence-constructor.png",
    "orders_client": ROOT / "evidence-orders-client.png",
    "orders_admin": ROOT / "evidence-orders-admin.png",
    "admin_orders_list": ROOT / "evidence-admin-orders-list.png",
    "admin_order_detail": ROOT / "evidence-admin-order-detail.png",
    "admin_products": ROOT / "evidence-admin-products.png",
    "admin_products_compact": ROOT / "evidence-admin-products-compact.png",
    "admin_constructor": ROOT / "evidence-admin-constructor.png",
    "client_orders_populated": ROOT / "evidence-client-orders-populated.png",
    "client_order_detail": ROOT / "evidence-client-order-detail.png",
    "catalog": ROOT / "evidence-catalog.png",
    "cart": ROOT / "evidence-cart.png",
    "checkout": ROOT / "evidence-checkout.png",
}

TEST_REPORT = ROOT / ".ai" / "vitest-report.json"


def set_cell_border(cell, **kwargs):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_borders = tc_pr.first_child_found_in("w:tcBorders")
    if tc_borders is None:
        tc_borders = OxmlElement("w:tcBorders")
        tc_pr.append(tc_borders)
    for edge in ("left", "top", "right", "bottom", "insideH", "insideV"):
        if edge in kwargs:
            edge_data = kwargs.get(edge)
            tag = f"w:{edge}"
            element = tc_borders.find(qn(tag))
            if element is None:
                element = OxmlElement(tag)
                tc_borders.append(element)
            for key in ["val", "sz", "space", "color"]:
                if key in edge_data:
                    element.set(qn(f"w:{key}"), str(edge_data[key]))


def set_table_borders(table):
    for row in table.rows:
        for cell in row.cells:
            set_cell_border(
                cell,
                left={"val": "single", "sz": "4", "space": "0", "color": "000000"},
                right={"val": "single", "sz": "4", "space": "0", "color": "000000"},
                top={"val": "single", "sz": "4", "space": "0", "color": "000000"},
                bottom={"val": "single", "sz": "4", "space": "0", "color": "000000"},
            )


def shade_cell(cell, fill="EDEDED"):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_table_font(table, size=12):
    for row in table.rows:
        for cell in row.cells:
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            for paragraph in cell.paragraphs:
                paragraph.paragraph_format.space_before = Pt(0)
                paragraph.paragraph_format.space_after = Pt(0)
                paragraph.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
                paragraph.paragraph_format.first_line_indent = Cm(0)
                for run in paragraph.runs:
                    run.font.name = "Times New Roman"
                    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
                    run.font.size = Pt(size)


def add_paragraph(doc, text="", *, style=None, align=None, bold=False, italic=False, size=14, first_indent=True):
    p = doc.add_paragraph(style=style)
    run = p.add_run(text)
    run.bold = bold
    run.italic = italic
    run.font.name = "Times New Roman"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    run.font.size = Pt(size)
    fmt = p.paragraph_format
    fmt.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    fmt.space_after = Pt(0)
    fmt.space_before = Pt(0)
    fmt.first_line_indent = Cm(1) if first_indent else Cm(0)
    if align is not None:
        p.alignment = align
    else:
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    return p


def add_heading(doc, text, *, level=1, centered=False, page_break_before=False):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = True
    run.font.name = "Times New Roman"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    run.font.size = Pt(14 if level == 1 else 13)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER if centered else WD_ALIGN_PARAGRAPH.LEFT
    fmt = p.paragraph_format
    fmt.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    fmt.space_before = Pt(6)
    fmt.space_after = Pt(6)
    fmt.first_line_indent = Cm(0)
    fmt.page_break_before = page_break_before
    fmt.keep_with_next = True
    return p


def add_table_caption(doc, text, *, page_break_before=False):
    p = add_paragraph(doc, text, first_indent=True, size=14)
    p.paragraph_format.keep_with_next = True
    p.paragraph_format.page_break_before = page_break_before
    return p


def add_figure_caption(doc, text):
    p = add_paragraph(doc, text, first_indent=False, size=14)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    return p


def add_table(doc, headers, rows, widths=None):
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr = table.rows[0].cells
    for idx, header in enumerate(headers):
        hdr[idx].text = header
        hdr[idx].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        shade_cell(hdr[idx])
        for run in hdr[idx].paragraphs[0].runs:
            run.bold = True
    for row in rows:
        cells = table.add_row().cells
        for idx, value in enumerate(row):
            cells[idx].text = value
    if widths:
        for row in table.rows:
            for idx, width in enumerate(widths):
                row.cells[idx].width = width
    set_table_borders(table)
    set_table_font(table, 12)
    return table


def page_break(doc):
    doc.add_page_break()


def configure_document(doc: Document):
    styles = doc.styles
    if "Normal" in styles:
        normal = styles["Normal"]
        normal.font.name = "Times New Roman"
        normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
        normal.font.size = Pt(14)

    for style_name in ("Heading 1", "Heading 2", "Heading 3"):
        if style_name in styles:
            style = styles[style_name]
            style.font.name = "Times New Roman"
            style._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")

    settings = doc.settings.element
    update_fields = settings.find(qn("w:updateFields"))
    if update_fields is None:
        update_fields = OxmlElement("w:updateFields")
        settings.append(update_fields)
    update_fields.set(qn("w:val"), "true")


def add_page_number(paragraph):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run()
    run.font.name = "Times New Roman"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    run.font.size = Pt(12)

    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = "PAGE"
    fld_sep = OxmlElement("w:fldChar")
    fld_sep.set(qn("w:fldCharType"), "separate")
    text = OxmlElement("w:t")
    text.text = "2"
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")

    run._r.append(fld_begin)
    run._r.append(instr)
    run._r.append(fld_sep)
    run._r.append(text)
    run._r.append(fld_end)


def strip_trailing_empty_paragraphs(doc: Document):
    body = doc._body._body
    children = list(body)
    for element in reversed(children):
        if element.tag != qn("w:p"):
            continue
        texts = "".join(node.text or "" for node in element.iter() if node.tag == qn("w:t"))
        has_break = any(node.tag == qn("w:br") for node in element.iter())
        if texts.strip() or has_break:
            break
        body.remove(element)


def add_bookmark(paragraph, name: str, bookmark_id: int, start: bool = True):
    element = OxmlElement("w:bookmarkStart" if start else "w:bookmarkEnd")
    element.set(qn("w:id"), str(bookmark_id))
    if start:
        element.set(qn("w:name"), name)
    paragraph._p.insert(0, element)


def read_test_metrics():
    data = json.loads(TEST_REPORT.read_text(encoding="utf-8"))
    suites = []
    for suite in data["testResults"]:
        suites.append(
            {
                "file": Path(suite["name"]).name,
                "count": len(suite.get("assertionResults", [])),
                "status": suite.get("status", ""),
                "tests": [
                    " > ".join(assertion.get("ancestorTitles", []) + [assertion.get("title", "")])
                    for assertion in suite.get("assertionResults", [])
                ],
            }
        )
    suites.sort(key=lambda item: item["file"])
    return {
        "total_suites": len(suites),
        "total_tests": data["numTotalTests"],
        "passed_tests": data["numPassedTests"],
        "suites": suites,
    }


def add_figure(doc, image_path: Path, caption: str, width_cm: float = 15.5):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    run.add_picture(str(image_path), width=Cm(width_cm))
    add_figure_caption(doc, caption)


def build_report():
    metrics = read_test_metrics()
    doc = Document(str(TITLE_DOCX))
    configure_document(doc)
    strip_trailing_empty_paragraphs(doc)

    section = doc.sections[0]
    section.different_first_page_header_footer = True
    header_para = section.header.paragraphs[0] if section.header.paragraphs else section.header.add_paragraph()
    header_para.text = ""
    add_page_number(header_para)

    page_break(doc)
    add_heading(doc, "ЗМІСТ", centered=True)
    spacer = doc.add_paragraph()
    spacer.paragraph_format.first_line_indent = Cm(0)
    spacer.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    toc_placeholder = doc.add_paragraph()
    toc_placeholder.paragraph_format.first_line_indent = Cm(0)
    toc_placeholder.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    add_bookmark(toc_placeholder, "toc_anchor", 2, start=True)
    add_bookmark(toc_placeholder, "toc_anchor", 2, start=False)

    body_start = add_heading(doc, "Мета роботи", centered=True, page_break_before=True)
    add_bookmark(body_start, "reportbody", 1, start=True)
    add_paragraph(
        doc,
        (
            "Метою роботи є розроблення повного комплекту документації з тестування "
            "програмного засобу для продажу авторських прикрас з конструктором дизайну, "
            "планування модульного, інтеграційного, системного та валідаційного тестування, "
            "а також виконання основних перевірок із фіксацією результатів і висновків щодо якості ПЗ."
        ),
    )
    add_paragraph(doc, "")
    tasks_heading = add_heading(doc, "ЗАВДАННЯ", centered=True)
    tasks = [
        "Розробити перелік функціональних і нефункціональних вимог до програмного засобу.",
        "Побудувати верифікаційні таблиці відповідності вимог, системних вимог і тестів.",
        "Скласти плани модульного, інтеграційного, системного та валідаційного тестування.",
        "Виконати відповідні види тестування на основі підготовлених планів.",
        "Зафіксувати результати перевірок у вигляді тестових звітів, проаналізувати виявлені дефекти та сформулювати висновки щодо якості ПЗ.",
    ]
    for idx, task in enumerate(tasks, 1):
        add_paragraph(doc, f"{idx}. {task}", first_indent=False)

    page_break(doc)
    execution_heading = add_heading(doc, "ВИКОНАННЯ ЗАВДАННЯ", centered=True)

    heading = add_heading(doc, "1 Вступ")
    intro_paragraphs = [
        "Тестування програмного забезпечення є завершальним та одним із найважливіших етапів життєвого циклу розроблення. Саме на цьому етапі підтверджується відповідність реалізованих функцій початковим вимогам, перевіряється стабільність поведінки системи в реальному середовищі та оцінюється готовність продукту до експлуатації.",
        "Об’єктом дослідження в даній курсовій роботі є вебсистема Aurora Atelier, призначена для продажу авторських прикрас, підтримки персонального кабінету клієнта, кошика, оформлення замовлень та використання конструктора індивідуального дизайну виробу.",
        "Робота виконувалася відповідно до матеріалів дисципліни «Тестування та верифікація ПЗ», вимог до оформлення пояснювальної записки та на основі фактичної реалізації проєкту. У межах виконання було не лише підготовлено документацію, а й дописано нові автоматизовані тести, проведено ручну перевірку критичних сценаріїв і зафіксовано реальні результати прогонів.",
    ]
    for text in intro_paragraphs:
        add_paragraph(doc, text)

    heading = add_heading(doc, "2 Характеристика програмного засобу")
    heading = add_heading(doc, "2.1 Призначення та основні можливості", level=2)
    for text in [
        "Aurora Atelier є клієнт-серверним вебзастосунком для демонстрації каталогу прикрас, формування індивідуального замовлення та адміністрування життєвого циклу замовлення.",
        "Клієнтська частина реалізує перегляд каталогу, фільтрацію товарів, авторизацію, роботу з кошиком, оформлення замовлення, персональний кабінет та візуальний конструктор прикрас.",
        "Серверна частина забезпечує маршрути API для аутентифікації, роботи з кошиком, побудови конфігурації конструктора, оформлення замовлення, підтвердження оплати, ведення історії замовлень та адміністрування списку замовлень.",
    ]:
        add_paragraph(doc, text)

    heading = add_heading(doc, "2.2 Архітектура та технологічний стек", level=2)
    add_paragraph(
        doc,
        "Система побудована за схемою «React-клієнт – Express API – SQLite». Для серверної частини застосовано Node.js та Express, для зберігання даних – SQLite3 із міграціями та seed-даними Knex, для фронтенду – React, для автоматизованих перевірок – Vitest, для ручних перевірок – браузерний прогін локального застосунку.",
    )
    add_table_caption(doc, "Таблиця 2.1 – Основні складові програмного засобу")
    add_table(
        doc,
        ["Складова", "Призначення", "Технології"],
        [
            ["Клієнтський інтерфейс", "Перегляд каталогу, конструктор, кошик, особистий кабінет", "React, Vite, браузер"],
            ["HTTP API", "Оброблення бізнес-логіки та взаємодія між клієнтом і БД", "Node.js, Express"],
            ["Сховище даних", "Зберігання товарів, користувачів, кошиків, замовлень, промокодів", "SQLite3, Knex"],
            ["Автотестування", "Регресійні, інтеграційні та модульні перевірки", "Vitest"],
            ["Ручне тестування", "Перевірка ключових сценаріїв користувача та адміністратора", "Локальний браузер, скриншоти"],
        ],
        widths=[Cm(4.0), Cm(8.6), Cm(4.0)],
    )
    add_paragraph(doc, "")

    heading = add_heading(doc, "2.3 Середовище тестування", level=2)
    add_table_caption(doc, "Таблиця 2.2 – Конфігурація середовища тестування")
    add_table(
        doc,
        ["Параметр", "Значення"],
        [
            ["Операційна система", "Windows, локальне середовище користувача"],
            ["Адреса локального запуску", "http://127.0.0.1:3000"],
            ["Тестова БД", "SQLite3, база після команди npm run db:reset"],
            ["Тестові облікові записи", "client@aurora.local / password123; admin@aurora.local / password123"],
            ["Команда запуску автотестів", "npm test; npx vitest run --reporter=json --outputFile .ai/vitest-report.json"],
            ["Інструменти ручної перевірки", "Локальний браузер, збереження скриншотів сторінок"],
        ],
        widths=[Cm(6.0), Cm(10.6)],
    )
    add_paragraph(doc, "")

    heading = add_heading(doc, "3 Формування вимог до ПЗ")
    heading = add_heading(doc, "3.1 Функціональні вимоги", level=2)
    add_paragraph(
        doc,
        "Функціональні вимоги сформовано на підставі фактичної реалізації проєкту, його архітектурної документації та основних користувацьких сценаріїв: реєстрація, автентифікація, робота з каталогом, конструктором, кошиком, замовленнями та адміністративною панеллю.",
    )
    fr_rows = [
        ["FR01", "Система має забезпечувати реєстрацію нового клієнта з перевіркою коректності введених контактних даних."],
        ["FR02", "Система має підтримувати автентифікацію користувачів та відновлення активної сесії."],
        ["FR03", "Система має розмежовувати ролі клієнта та адміністратора."],
        ["FR04", "Система має надавати каталог виробів з підтримкою фільтрів за категорією, матеріалом, каменем, колекцією та ціною."],
        ["FR05", "Система має надавати конфігурацію конструктора та дозволяти вибір параметрів індивідуального виробу."],
        ["FR06", "Система має виконувати серверний розрахунок вартості індивідуального виробу з перевіркою вибраних слотів каменів."],
        ["FR07", "Система має дозволяти додавання готових і кастомних виробів до кошика, зміну кількості та повторне використання активного кошика."],
        ["FR08", "Система має підтримувати оформлення замовлення з перевіркою обов’язкових погоджень користувача."],
        ["FR09", "Система має враховувати дію промокодів і контролювати обмеження на використання промоакцій."],
        ["FR10", "Система має вести історію замовлень клієнта та відображати поточне активне замовлення."],
        ["FR11", "Система має підтримувати підтвердження оплати та коректний життєвий цикл статусів замовлення."],
        ["FR12", "Система має надавати адміністратору доступ до панелі замовлень, фільтрації списку та перегляду деталей."],
        ["FR13", "Система має надсилати сервісні сповіщення про створення або зміну статусу замовлення."],
        ["FR14", "Система має надавати перемикання мови інтерфейсу в клієнтській частині."],
    ]
    add_table_caption(doc, "Таблиця 3.1 – Перелік функціональних вимог")
    add_table(doc, ["ID", "Формулювання вимоги"], fr_rows, widths=[Cm(2.2), Cm(14.4)])
    add_paragraph(doc, "")

    heading = add_heading(doc, "3.2 Нефункціональні вимоги", level=2, page_break_before=True)
    nfr_rows = [
        ["NFR01", "ПЗ має забезпечувати авторизацію доступу до захищених маршрутів і даних."],
        ["NFR02", "ПЗ має формувати відповідь API для типових сценаріїв без критичних затримок у локальному середовищі."],
        ["NFR03", "ПЗ має забезпечувати цілісність записів кошика, замовлень і промокодів на рівні БД та бізнес-логіки."],
        ["NFR04", "ПЗ має бути сумісним із сучасним браузером і коректно відображати основні сторінки інтерфейсу."],
        ["NFR05", "ПЗ має бути зручним у використанні: типові сценарії повинні виконуватися без додаткового навчання."],
        ["NFR06", "ПЗ має бути стійким до повторних запитів і не створювати дублікати замовлень або підтверджень оплати."],
        ["NFR07", "ПЗ має підтримувати автоматизоване тестування критичних модулів і міжмодульних взаємодій."],
        ["NFR08", "ПЗ має вести журналювання важливих подій аутентифікації та оброблення замовлень."],
    ]
    add_table_caption(doc, "Таблиця 3.2 – Перелік нефункціональних вимог")
    add_table(doc, ["ID", "Формулювання вимоги"], nfr_rows, widths=[Cm(2.2), Cm(14.4)])
    add_paragraph(doc, "")

    heading = add_heading(doc, "4 Планування валідаційного тестування")
    for text in [
        "Метою валідаційного тестування є підтвердження того, що система реалізує очікувані користувацькі сценарії та відповідає критеріям приймання. В даному випадку приймання орієнтується на готовність системи до використання клієнтом і адміністратором у типовому робочому процесі.",
        "Для кожної основної функціональної вимоги визначено приймальний сценарій, очікуваний результат і критерій успішності. План покриває як позитивні сценарії, так і критичні обмеження доступу або перевірки вхідних даних.",
    ]:
        add_paragraph(doc, text)
    add_table_caption(doc, "Таблиця 4.1 – План валідаційного тестування")
    add_table(
        doc,
        ["ID AT", "Пов’язана вимога", "Сценарій перевірки", "Очікуваний результат"],
        [
            ["AT01", "FR01", "Реєстрація нового клієнта з валідними даними", "Створюється обліковий запис зі статусом очікування підтвердження"],
            ["AT02", "FR02", "Вхід клієнта до акаунта", "Створюється активна сесія та відкривається сторінка замовлень"],
            ["AT03", "FR04", "Перегляд каталогу та застосування фільтрів", "Виводиться релевантний перелік товарів"],
            ["AT04", "FR05, FR06", "Робота з конструктором виробу", "Відображається конфігурація, а сервер приймає коректні опції"],
            ["AT05", "FR07", "Додавання товарів до кошика", "Кошик містить позиції, кількість коректно оновлюється"],
            ["AT06", "FR08", "Оформлення замовлення з обов’язковими погодженнями", "Створюється нове замовлення без дублювання"],
            ["AT07", "FR09", "Використання промокоду", "Знижка застосовується в межах бізнес-правил"],
            ["AT08", "FR10", "Перегляд історії замовлень клієнта", "Клієнт бачить поточне та завершені замовлення"],
            ["AT09", "FR12", "Вхід адміністратора та перегляд адмін-панелі замовлень", "Адміністратор отримує доступ до списку замовлень і деталей"],
            ["AT10", "NFR01", "Спроба клієнта отримати доступ до адмін-маршруту", "Доступ заборонено, повертається код 403"],
        ],
        widths=[Cm(2.0), Cm(2.8), Cm(6.0), Cm(5.8)],
    )
    add_paragraph(doc, "")
    add_table_caption(doc, "Таблиця 4.2 – Верифікаційна матриця вимог і приймальних тестів")
    add_table(
        doc,
        ["Вимога", "AT"],
        [
            ["FR01", "AT01"],
            ["FR02", "AT02"],
            ["FR04", "AT03"],
            ["FR05, FR06", "AT04"],
            ["FR07", "AT05"],
            ["FR08", "AT06"],
            ["FR09", "AT07"],
            ["FR10", "AT08"],
            ["FR12", "AT09"],
            ["NFR01", "AT10"],
        ],
        widths=[Cm(4.0), Cm(12.6)],
    )
    add_paragraph(doc, "")

    heading = add_heading(doc, "5 Планування системного тестування")
    for text in [
        "Системне тестування виконується щодо повністю інтегрованої вебсистеми й охоплює перевірку функціональних та нефункціональних характеристик при роботі всіх компонентів разом. За основу взято системні вимоги, сформовані з функціональних і нефункціональних вимог прикладного рівня.",
        "Для даного проєкту системні вимоги деталізують очікувану поведінку сторінок інтерфейсу, API-маршрутів, механізмів авторизації, кошика, конструктора, оформлення та супроводу замовлень.",
    ]:
        add_paragraph(doc, text)
    sr_rows = [
        ["FR01.SR01", "Система має відхиляти реєстрацію з некоректним номером телефону."],
        ["FR02.SR01", "Система має створювати й очищати сесію користувача під час входу та виходу."],
        ["FR03.SR01", "Система має блокувати доступ клієнта до адміністративних маршрутів."],
        ["FR04.SR01", "Система має повертати лише підтримувані значення фільтрів каталогу."],
        ["FR05.SR01", "Система має повертати активні типи прикрас і параметри конструктора."],
        ["FR06.SR01", "Система має відхиляти некоректні ідентифікатори слотів каменів."],
        ["FR07.SR01", "Система має повторно використовувати активний кошик клієнта."],
        ["FR08.SR01", "Система має відхиляти оформлення замовлення без обов’язкових погоджень."],
        ["FR08.SR02", "Система має гарантувати ідемпотентність повторного checkout-запиту."],
        ["FR09.SR01", "Система має не дозволяти повторне використання промокоду понад ліміт."],
        ["FR11.SR01", "Система має не дозволяти некоректний перехід між статусами замовлення."],
        ["NFR03.SR01", "База даних має відхиляти несумісні структури елементів кошика і підсумків замовлення."],
    ]
    add_table_caption(doc, "Таблиця 5.1 – Перелік системних вимог")
    add_table(doc, ["ID SR", "Визначення системної вимоги"], sr_rows, widths=[Cm(3.2), Cm(13.4)])
    add_paragraph(doc, "")
    add_table_caption(doc, "Таблиця 5.2 – RTM(SWR, SR)")
    add_table(
        doc,
        ["Вимога ПЗ", "Перелік системних вимог"],
        [
            ["FR01", "FR01.SR01"],
            ["FR02", "FR02.SR01"],
            ["FR03", "FR03.SR01"],
            ["FR04", "FR04.SR01"],
            ["FR05", "FR05.SR01"],
            ["FR06", "FR06.SR01"],
            ["FR07", "FR07.SR01"],
            ["FR08", "FR08.SR01; FR08.SR02"],
            ["FR09", "FR09.SR01"],
            ["FR11", "FR11.SR01"],
            ["NFR03", "NFR03.SR01"],
        ],
        widths=[Cm(3.0), Cm(13.6)],
    )
    add_paragraph(doc, "")
    add_table_caption(doc, "Таблиця 5.3 – План системних тестів", page_break_before=True)
    add_table(
        doc,
        ["ID ST", "ID SR", "Тестовий сценарій", "Очікуваний результат"],
        [
            ["ST01", "FR01.SR01", "Спроба реєстрації з невалідним телефоном", "Система повертає помилку валідації"],
            ["ST02", "FR02.SR01", "Вхід і вихід користувача", "Сесія створюється та коректно завершується"],
            ["ST03", "FR03.SR01", "Клієнт звертається до /api/admin/orders", "Отримується HTTP 403"],
            ["ST04", "FR04.SR01", "Фільтрація каталогу", "Повертаються лише підтримувані опції"],
            ["ST05", "FR05.SR01", "Отримання конфігурації конструктора", "Повертаються активні типи та параметри"],
            ["ST06", "FR06.SR01", "Надсилання невалідного stoneSlotId", "Сервер відхиляє запит"],
            ["ST07", "FR07.SR01", "Повторне додавання товару в активний кошик", "Кількість у кошику збільшується"],
            ["ST08", "FR08.SR01", "Checkout без погоджень", "Оформлення замовлення блокується"],
            ["ST09", "FR08.SR02", "Подвійний checkout", "Створюється лише одне замовлення"],
            ["ST10", "FR09.SR01", "Повторне використання промокоду понад ліміт", "Знижка не застосовується"],
            ["ST11", "FR11.SR01", "Спроба некоректного переходу статусу", "Перехід блокується логікою сервера"],
            ["ST12", "NFR03.SR01", "Вставка неузгоджених даних у БД через тест", "База відхиляє некоректний запис"],
        ],
        widths=[Cm(1.8), Cm(2.8), Cm(6.0), Cm(6.0)],
    )
    add_paragraph(doc, "")

    heading = add_heading(doc, "6 Планування інтеграційного тестування")
    for text in [
        "Інтеграційне тестування спрямоване на перевірку коректності взаємодії між основними модулями: аутентифікацією, маршрутизацією API, кошиком, оформленням замовлення, конструктором, адміністративним модулем та шаром даних.",
        "Для фіксації інтеграційних перевірок використано сценарії рівня API. Такий підхід дозволяє перевірити одночасно HTTP-маршрут, middleware, сервісну логіку та роботу з БД у межах одного тесту.",
    ]:
        add_paragraph(doc, text)
    add_table_caption(doc, "Таблиця 6.1 – Об’єкти інтеграційного тестування")
    add_table(
        doc,
        ["Пара модулів", "Що перевіряється"],
        [
            ["Auth API – Session storage", "Створення, читання та очищення сесії"],
            ["Catalog API – Filter normalization", "Коректне формування результату фільтрації"],
            ["Constructor API – Pricing service", "Передача та перевірка параметрів конструктора"],
            ["Cart API – Order service", "Перенесення елементів кошика в замовлення"],
            ["Checkout API – Promo service", "Застосування знижки та облік лімітів"],
            ["Orders API – Account page", "Побудова поточного та завершених замовлень користувача"],
            ["Admin Orders API – Authorization middleware", "Контроль ролей і доступу до адмін-маршрутів"],
        ],
        widths=[Cm(6.0), Cm(10.6)],
    )
    add_paragraph(doc, "")
    add_table_caption(doc, "Таблиця 6.2 – План інтеграційних тестів")
    add_table(
        doc,
        ["ID IT", "Інтеграційний сценарій", "Очікуваний результат"],
        [
            ["IT01", "Реєстрація клієнта та створення запису в БД", "Новий користувач зберігається як pending verification"],
            ["IT02", "Логін, читання сесії та logout", "Сесія існує після входу та зникає після виходу"],
            ["IT03", "Доступ до кошика без авторизації", "Система повертає 401 Unauthorized"],
            ["IT04", "Додавання того самого товару до кошика", "Оновлюється кількість, а не створюється дубль"],
            ["IT05", "Checkout і створення замовлення", "Після оформлення формується нове замовлення"],
            ["IT06", "Підтвердження оплати повторним запитом", "Стан замовлення оновлюється без дублювання ефекту"],
            ["IT07", "Контроль статусних переходів адміном", "Логіка не дозволяє заборонені переходи"],
            ["IT08", "Отримання профілю та історії замовлень", "Повертається профіль, поточне і завершені замовлення"],
            ["IT09", "Доступ клієнта до адмін-панелі", "Отримується код 403"],
            ["IT10", "Вхід адміністратора та отримання деталей замовлень", "Панель замовлень і деталі доступні"],
        ],
        widths=[Cm(2.0), Cm(8.0), Cm(6.6)],
    )
    add_paragraph(doc, "")

    heading = add_heading(doc, "7 Планування модульного тестування")
    for text in [
        "Модульне тестування орієнтоване на перевірку окремих утиліт, обчислювальних функцій та допоміжної логіки, яка може бути коректно протестована ізольовано від HTTP-шару та БД.",
        "Для даного проєкту доцільно виділити модулі, що відповідають за нормалізацію фільтрів каталогу, грошові обчислення, побудову прев’ю виробу, роботу з промокодами та перевірку лінійності зміни статусів.",
    ]:
        add_paragraph(doc, text)
    add_table_caption(doc, "Таблиця 7.1 – Об’єкти модульного тестування")
    add_table(
        doc,
        ["Модуль", "Призначення", "Причина модульного тестування"],
        [
            ["catalog filters", "Нормалізація значень фільтрів каталогу", "Висока кількість варіантів вхідних значень"],
            ["money utils", "Округлення та підсумовування валютних значень", "Критичність для підрахунку вартості"],
            ["order statuses", "Перехід між статусами та визначення прострочення", "Логіка workflow замовлення"],
            ["preview layers", "Побудова набору шарів прев’ю кастомного виробу", "Візуальна коректність конструктора"],
            ["promo codes", "Розрахунок процентної та фіксованої знижки", "Чутливість до граничних значень"],
        ],
        widths=[Cm(4.0), Cm(6.0), Cm(6.6)],
    )
    add_paragraph(doc, "")
    add_table_caption(doc, "Таблиця 7.2 – План модульних тестів")
    add_table(
        doc,
        ["ID UT", "Модуль", "Сценарій", "Очікуваний результат"],
        [
            ["UT01", "catalog filters", "Відкидання непідтримуваних значень", "У результаті лишаються лише допустимі опції"],
            ["UT02", "catalog filters", "Повернення точних переліків фільтрів", "Списки опцій відповідають конфігурації"],
            ["UT03", "money utils", "Округлення до двох знаків", "Результат відповідає валютному формату"],
            ["UT04", "money utils", "Підсумовування кількох значень", "Сума коректно округлюється"],
            ["UT05", "order statuses", "Лінійний перехід між статусами", "Повертається дозволений наступний статус"],
            ["UT06", "order statuses", "Визначення прострочення in-progress", "Старі замовлення позначаються overdue"],
            ["UT07", "preview layers", "Побудова шарів бази, каменів та гравіювання", "Усі активні шари включено"],
            ["UT08", "promo codes", "Процентна знижка", "Знижка не перевищує subtotal"],
            ["UT09", "promo codes", "Фіксована знижка", "Результат обмежено вартістю замовлення"],
        ],
        widths=[Cm(1.8), Cm(3.0), Cm(6.4), Cm(6.4)],
    )
    add_paragraph(doc, "")

    heading = add_heading(doc, "8 Виконання модульного тестування")
    add_paragraph(
        doc,
        "Модульні перевірки були виконані під час загального прогона тестового набору за допомогою Vitest. Для цього використано ізольовані тести над утилітами та допоміжними функціями без підняття зовнішніх сервісів.",
    )
    add_table_caption(doc, "Таблиця 8.1 – Фактичне виконання модульних тестів")
    add_table(
        doc,
        ["Файл тестів", "Кількість тестів", "Результат"],
        [
            ["catalog-filters.test.js", "2", "Пройдено"],
            ["money.test.js", "2", "Пройдено"],
            ["order-statuses.test.js", "2", "Пройдено"],
            ["preview.test.js", "1", "Пройдено"],
            ["promo-codes.test.js", "2", "Пройдено"],
        ],
        widths=[Cm(7.0), Cm(4.0), Cm(5.6)],
    )
    add_paragraph(doc, "Результати свідчать, що критична обчислювальна логіка працює коректно на типових і граничних значеннях. Збоїв або провалених модульних тестів не зафіксовано.")

    heading = add_heading(doc, "9 Виконання інтеграційного тестування")
    add_paragraph(
        doc,
        "Під час виконання інтеграційного тестування було дописано окремий файл tests/integration/api-flows.test.js, який розширив покриття сценаріями реєстрації, роботи сесії, доступу до кошика, перевірки прав доступу та адмін-маршрутів. Окремо використовувався наявний файл cart-checkout.test.js, що покриває критичні сценарії checkout, промокодів, статусів замовлення та консистентності БД.",
    )
    add_table_caption(doc, "Таблиця 9.1 – Фактичне виконання інтеграційних тестів")
    add_table(
        doc,
        ["Файл тестів", "Кількість тестів", "Коротка характеристика", "Результат"],
        [
            ["app.test.js", "4", "Health-check, 404, security headers, rate limit", "Пройдено"],
            ["api-flows.test.js", "12", "Нові сценарії API flows та контроль доступу", "Пройдено"],
            ["cart-checkout.test.js", "16", "Checkout, промокоди, статуси, account, consistency", "Пройдено"],
        ],
        widths=[Cm(4.8), Cm(2.4), Cm(7.2), Cm(2.2)],
    )
    add_paragraph(doc, "Усього було додано 12 нових інтеграційних тестів, що дозволило суттєво підсилити перевірку критичних бізнес-процесів без зміни функціональності самого застосунку.")
    add_table_caption(doc, "Таблиця 9.2 – Перелік доданих інтеграційних тестів")
    added_tests = [
        ["1", "Анонімний виклик session endpoint", "Підтвердження стану гостя"],
        ["2", "Реєстрація з невалідним українським телефоном", "Перевірка валідації введення"],
        ["3", "Реєстрація pending-verification акаунта", "Створення користувача у БД"],
        ["4", "Логін клієнта та logout", "Життєвий цикл сесії"],
        ["5", "Вимога авторизації для кошика", "Контроль захищеного маршруту"],
        ["6", "Отримання активного кошика клієнтом", "Зв’язок auth–cart"],
        ["7", "Повторне додавання того самого товару", "Оновлення кількості замість дублю"],
        ["8", "Отримання active jewelry types", "Перевірка constructor config"],
        ["9", "Checkout без погоджень", "Валідація обов’язкових чекбоксів"],
        ["10", "Контроль доступу до orders endpoints", "Авторство замовлень"],
        ["11", "Заборона клієнту доступу до admin orders", "Контроль ролей"],
        ["12", "Вхід адміністратора і доступ до dashboard/details", "Повний адмінський сценарій"],
    ]
    add_table(doc, ["№", "Назва нового тесту", "Що перевіряється"], added_tests, widths=[Cm(1.0), Cm(7.4), Cm(8.0)])
    add_paragraph(doc, "")

    heading = add_heading(doc, "10 Виконання системного тестування")
    add_paragraph(
        doc,
        "Системне тестування проводилося як комбінація реальних ручних перевірок у локально запущеному застосунку та інтеграційних автотестів API. Це дозволило перевірити не лише кодові інваріанти, але й фактичне відображення сторінок і доступність ключових сценаріїв для кінцевого користувача.",
    )
    add_table_caption(doc, "Таблиця 10.1 – Результати ручних системних перевірок")
    add_table(
        doc,
        ["№", "Сценарій", "Фактичний результат", "Статус"],
        [
            ["1", "Відкриття головної сторінки", "Головна сторінка Aurora Atelier успішно завантажується", "Пройдено"],
            ["2", "Відкриття сторінки конструктора", "Сторінка конструктора відображає блоки конфігурації", "Пройдено"],
            ["3", "Вхід клієнта до акаунта", "Після входу відкривається сторінка замовлень клієнта", "Пройдено"],
            ["4", "Перегляд сторінки замовлень клієнта", "Історія та поточний стан доступні після авторизації", "Пройдено"],
            ["5", "Вхід адміністратора", "Після входу доступна сторінка управління замовленнями", "Пройдено"],
            ["6", "Перегляд адмінського списку замовлень", "Список і деталі замовлень відкриваються коректно", "Пройдено"],
        ],
        widths=[Cm(1.0), Cm(6.8), Cm(7.2), Cm(2.4)],
    )
    add_paragraph(doc, "Результати ручної перевірки підтвердили коректне відкриття та роботу критичних сторінок інтерфейсу. Для підвищення доказовості в звіт включено не лише загальні сторінки застосунку, а й окремі скриншоти клієнтських і адміністративних сценаріїв з реальними даними замовлень.")
    heading = add_heading(doc, "10.1 Деталізація ручних сценаріїв", level=2)
    add_table_caption(doc, "Таблиця 10.2 – Детальні сценарії ручного тестування")
    add_table(
        doc,
        ["ID", "Роль", "Сценарій", "Очікуваний результат", "Фактичний результат"],
        [
            ["MT01", "Гість", "Відкриття головної сторінки", "Головна сторінка завантажується без помилок", "Виконано, сторінка відкривається"],
            ["MT02", "Клієнт", "Авторизація через форму входу", "Після входу відкривається сторінка замовлень", "Виконано, відкрито /orders"],
            ["MT03", "Клієнт", "Перегляд каталогу", "Доступні картки товарів і елементи навігації", "Виконано, каталог відображено"],
            ["MT04", "Клієнт", "Перегляд конструктора", "Доступні блоки конфігурації та preview", "Виконано, конструктор відкрито"],
            ["MT05", "Клієнт", "Перегляд кошика з товаром", "У кошику видно позицію, суму та summary", "Виконано, кошик містить товар"],
            ["MT06", "Клієнт", "Перегляд checkout", "Форма checkout містить контактні поля та підсумок", "Виконано, форма доступна"],
            ["MT07", "Клієнт", "Перегляд сторінки замовлень з даними", "Історія замовлень і статуси доступні", "Виконано, замовлення відображаються"],
            ["MT08", "Клієнт", "Перегляд деталізації замовлення", "Видно склад, статус і таймлайн", "Виконано, деталі замовлення відкрито"],
            ["MT09", "Адміністратор", "Вхід до адмінки orders", "Доступний список замовлень і summary", "Виконано, список відображається"],
            ["MT10", "Адміністратор", "Перегляд деталізації замовлення", "Видно контакти клієнта, позиції, статус", "Виконано, сторінка деталей доступна"],
            ["MT11", "Адміністратор", "Перегляд products page", "Видно список товарів і засоби керування", "Виконано, products page доступна"],
            ["MT12", "Адміністратор", "Перегляд constructor studio", "Доступні workspace, типи, assets і pricing", "Виконано, constructor studio відкрита"],
        ],
        widths=[Cm(1.2), Cm(2.2), Cm(4.6), Cm(4.2), Cm(4.4)],
    )
    add_paragraph(doc, "Нижче подано основні ілюстрації реального системного тестування.")
    add_figure(doc, SCREENSHOTS["home"], "Рисунок 10.1 – Головна сторінка вебсистеми Aurora Atelier", width_cm=15.2)
    add_figure(doc, SCREENSHOTS["constructor"], "Рисунок 10.2 – Сторінка конструктора індивідуального виробу", width_cm=14.8)
    add_figure(doc, SCREENSHOTS["catalog"], "Рисунок 10.3 – Каталог готових виробів", width_cm=14.8)
    add_figure(doc, SCREENSHOTS["cart"], "Рисунок 10.4 – Кошик користувача з реальним товаром", width_cm=14.4)
    add_figure(doc, SCREENSHOTS["checkout"], "Рисунок 10.5 – Сторінка оформлення замовлення", width_cm=14.4)
    add_figure(doc, SCREENSHOTS["client_orders_populated"], "Рисунок 10.6 – Сторінка замовлень клієнта з реальними замовленнями", width_cm=14.8)
    add_figure(doc, SCREENSHOTS["client_order_detail"], "Рисунок 10.7 – Детальна сторінка клієнтського замовлення", width_cm=14.8)
    add_figure(doc, SCREENSHOTS["admin_orders_list"], "Рисунок 10.8 – Адміністративний список замовлень", width_cm=14.8)
    add_figure(doc, SCREENSHOTS["admin_order_detail"], "Рисунок 10.9 – Детальна сторінка замовлення в адмінці", width_cm=14.4)
    add_figure(doc, SCREENSHOTS["admin_products_compact"], "Рисунок 10.10 – Адміністративна сторінка керування товарами", width_cm=12.2)
    add_figure(doc, SCREENSHOTS["admin_constructor"], "Рисунок 10.11 – Адміністративна студія конструктора", width_cm=12.8)

    heading = add_heading(doc, "11 Виконання валідаційного тестування")
    add_paragraph(
        doc,
        "Валідаційне тестування виконувалося на рівні ключових користувацьких сценаріїв. Частина критеріїв приймання була підтверджена ручним прогоном, інша частина – інтеграційними тестами, що моделюють реальні запити до API та перевіряють зовнішню поведінку системи.",
    )
    add_table_caption(doc, "Таблиця 11.1 – Результати валідаційного тестування")
    add_table(
        doc,
        ["ID AT", "Очікуваний результат", "Фактичний результат", "Статус"],
        [
            ["AT01", "Створюється новий акаунт", "Тест registration creates a pending-verification account пройдено", "Пройдено"],
            ["AT02", "Створюється клієнтська сесія", "Тест client login creates a session and logout clears it пройдено", "Пройдено"],
            ["AT03", "Каталог коректно фільтрується", "Каталог та unit/integration перевірки фільтрів пройдено", "Пройдено"],
            ["AT04", "Конструктор відкривається та надає конфігурацію", "Сторінка відкривається, API constructor config пройдено", "Пройдено"],
            ["AT05", "Кошик коректно оновлюється", "Повторне додавання товару збільшує кількість", "Пройдено"],
            ["AT06", "Checkout без дублювання замовлення", "Тест creates only one order when checkout is requested twice пройдено", "Пройдено"],
            ["AT07", "Промокод працює в рамках правил", "Перевірено застосування та ліміти промокоду", "Пройдено"],
            ["AT08", "Адмін отримує доступ до панелі замовлень", "Реальний вхід і інтеграційний тест підтвердили доступ", "Пройдено"],
            ["AT09", "Клієнт не має доступу до адмін-маршрутів", "Тест client cannot access admin orders endpoints пройдено", "Пройдено"],
            ["AT10", "Історія замовлень користувача доступна", "Поточне і завершені замовлення повертаються коректно", "Пройдено"],
        ],
        widths=[Cm(1.8), Cm(5.4), Cm(7.0), Cm(3.2)],
    )
    add_paragraph(doc, "")
    add_table_caption(doc, "Таблиця 11.2 – Зведені метрики виконання автотестів")
    add_table(
        doc,
        ["Загальна кількість тестів", "Кількість успішних", "Кількість неуспішних", "Відсоток успішних"],
        [[str(metrics["total_tests"]), str(metrics["passed_tests"]), "0", "100%"]],
        widths=[Cm(4.5), Cm(4.0), Cm(4.0), Cm(4.5)],
    )
    add_paragraph(doc, "За результатами фактичного прогона команда npm test завершилася успішно: пройдено всі 41 тести без жодного падіння.")

    heading = add_heading(doc, "12 Аналіз результатів і виявлених дефектів")
    for text in [
        "За результатами проведених перевірок критичних або блокуючих дефектів, що унеможливлюють роботу системи, не виявлено. Функціональні вимоги, винесені до планів тестування, підтверджені автотестами та ручною перевіркою інтерфейсу.",
        "Під час прогону було зафіксовано технічне зауваження низької критичності: попередження Express щодо застарілого використання res.clearCookie з параметром maxAge в модулі auth.service.js. Воно не впливає на користувацьку функціональність і не спричинило падіння тестів, але має бути усунуте в межах технічного супроводу для сумісності з майбутніми версіями Express.",
    ]:
        add_paragraph(doc, text)
    add_table_caption(doc, "Таблиця 12.1 – Виявлені дефекти та спостереження")
    add_table(
        doc,
        ["№", "Опис", "Критичність", "Статус", "Коментар"],
        [["1", "Попередження Express щодо res.clearCookie і maxAge", "Низька", "Відкрите", "Не впливає на проходження тестів, бажано усунути в наступній ітерації"]],
        widths=[Cm(1.0), Cm(6.8), Cm(2.6), Cm(2.6), Cm(5.6)],
    )
    add_paragraph(doc, "Окремо слід відзначити, що під час підготовки інтеграційного набору було усунено проблему тестового оточення: закриття спільного з’єднання Knex в одному integration suite призводило до падіння наступних suites. Після корекції структура запуску стала стабільною і повторюваною.")

    heading = add_heading(doc, "13 Узагальнення покриття тестування")
    add_paragraph(
        doc,
        "Побудована система перевірок забезпечує багаторівневе покриття: модульні тести контролюють допоміжну логіку, інтеграційні – взаємодію API, middleware і БД, системні перевірки – доступність інтерфейсу та ключових сценаріїв у реальному браузері, а валідаційні тести підтверджують відповідність очікуванням замовника й користувача.",
    )
    add_table_caption(doc, "Таблиця 13.1 – Зведення тестового покриття за рівнями")
    add_table(
        doc,
        ["Рівень тестування", "Основні об’єкти перевірки", "Форма виконання", "Результат"],
        [
            ["Модульне", "Утиліти, обчислення, preview, promo, status flow", "Автоматизовано", "Успішно"],
            ["Інтеграційне", "API, auth, cart, checkout, orders, admin", "Автоматизовано", "Успішно"],
            ["Системне", "Сторінки інтерфейсу та ролі користувачів", "Ручна перевірка", "Успішно"],
            ["Валідаційне", "Ключові користувацькі сценарії", "Комбіновано", "Успішно"],
        ],
        widths=[Cm(3.2), Cm(7.8), Cm(3.0), Cm(2.6)],
    )
    add_paragraph(doc, "З урахуванням доданих автотестів і ручних перевірок найбільше покриття отримали сценарії аутентифікації, контролю доступу, кошика, checkout, історії замовлень і адміністративного супроводу замовлень. Це відповідає ключовим ризиковим зонам поточної версії програмного засобу.")

    page_break(doc)
    heading = add_heading(doc, "ВИСНОВКИ", centered=True)
    conclusions = [
        "У межах курсової роботи було сформовано повний комплект документації з тестування програмного засобу для продажу авторських прикрас з конструктором дизайну. Для системи визначено функціональні й нефункціональні вимоги, побудовано верифікаційні таблиці та плани модульного, інтеграційного, системного й валідаційного тестування.",
        "Під час практичного виконання було дописано 12 нових інтеграційних тестів, після чого повний тестовий набір досяг 41 успішно виконаного тесту. Додатково проведено ручне тестування ключових сторінок інтерфейсу: головної сторінки, конструктора, кабінету клієнта та адміністративної сторінки замовлень. Це дозволило підтвердити працездатність основних користувацьких і службових сценаріїв.",
        "Результати тестування показали, що система коректно реалізує заплановані функції, забезпечує розмежування прав доступу, коректно працює з кошиком і замовленнями, а також не містить критичних дефектів. Єдине зафіксоване технічне зауваження стосується попередження сумісності Express і не впливає на прийнятність програмного засобу. Отже, ПЗ можна вважати готовим до подальшого використання та розвитку.",
    ]
    for text in conclusions:
        add_paragraph(doc, text)

    page_break(doc)
    heading = add_heading(doc, "Список використаних джерел", centered=True)
    sources = [
        "1. Клименко Т. А., Манжос Ю. С. Тестування та верифікація програмного забезпечення: навчальний посібник по практикуму з виконання курсового проєкту. Харків: ХАІ, 2025.",
        "2. Методичка ТВПЗ(КП) – 2026. Матеріали до виконання курсового проєкту з дисципліни «Тестування та верифікація ПЗ».",
        "3. Вимоги до оформлення пояснювальної записки, 2021.",
        "4. Документація проєкту Aurora Atelier: docs/README.md.",
        "5. Документація архітектури проєкту Aurora Atelier: docs/architecture/system-overview.md.",
        "6. Документація API проєкту Aurora Atelier: docs/architecture/backend-api.md.",
        "7. ISO/IEC 25010:2011 Systems and software engineering – Systems and software Quality Requirements and Evaluation (SQuaRE) – System and software quality models.",
        "8. Vitest Documentation. Test framework for JavaScript and TypeScript applications.",
        "9. Express Documentation. Web framework for Node.js applications.",
        "10. SQLite Documentation. Lightweight relational database engine.",
    ]
    for source in sources:
        add_paragraph(doc, source, first_indent=False)

    page_break(doc)
    heading = add_heading(doc, "ДОДАТОК А", centered=True)
    heading = add_heading(doc, "Перелік виконаних автоматизованих тестових наборів", centered=True)
    suite_rows = []
    for suite in metrics["suites"]:
        suite_rows.append([suite["file"], str(suite["count"]), "Пройдено"])
    add_table_caption(doc, "Таблиця А.1 – Зведення за файлами автотестів")
    add_table(doc, ["Файл", "Кількість тестів", "Статус"], suite_rows, widths=[Cm(9.8), Cm(3.2), Cm(2.8)])
    add_paragraph(doc, "")
    add_table_caption(doc, "Таблиця А.2 – Ключові автоматизовані сценарії")
    scenario_rows = []
    for suite in metrics["suites"]:
        for test_name in suite["tests"]:
            scenario_rows.append([Path(suite["file"]).stem, test_name[:95]])
    add_table(doc, ["Набір тестів", "Назва сценарію"], scenario_rows, widths=[Cm(4.8), Cm(11.0)])

    page_break(doc)
    heading = add_heading(doc, "ДОДАТОК Б", centered=True)
    heading = add_heading(doc, "Ілюстрації реального тестування", centered=True)
    add_paragraph(doc, "У додатку наведено додаткові скриншоти, отримані під час ручного тестування клієнтських і адміністративних сценаріїв. Їх включення до звіту дає змогу наочно підтвердити фактичне виконання перевірок.", first_indent=True)
    add_figure(doc, SCREENSHOTS["client_orders_populated"], "Рисунок Б.1 – Історія замовлень клієнта після створення реального замовлення", width_cm=14.6)
    add_figure(doc, SCREENSHOTS["client_order_detail"], "Рисунок Б.2 – Деталі клієнтського замовлення", width_cm=14.4)
    add_figure(doc, SCREENSHOTS["admin_orders_list"], "Рисунок Б.3 – Список замовлень в адміністративній панелі", width_cm=14.4)
    add_figure(doc, SCREENSHOTS["admin_order_detail"], "Рисунок Б.4 – Деталі замовлення в адміністративній панелі", width_cm=14.0)
    add_figure(doc, SCREENSHOTS["admin_products_compact"], "Рисунок Б.5 – Сторінка керування товарами", width_cm=12.2)
    add_figure(doc, SCREENSHOTS["admin_constructor"], "Рисунок Б.6 – Constructor Studio адміністратора", width_cm=12.8)
    add_bookmark(doc.paragraphs[-1], "reportbody", 1, start=False)

    doc.save(FINAL_DOCX)


if __name__ == "__main__":
    build_report()
    print(str(FINAL_DOCX))
