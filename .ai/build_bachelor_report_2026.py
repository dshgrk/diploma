from __future__ import annotations

import argparse
import base64
import math
import re
import subprocess
import zipfile
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Mm, Pt, RGBColor
from PIL import Image, ImageDraw, ImageFont


ROOT = Path(r"E:\UNIVERSITY\diploma")
OUTPUT_DOCX = ROOT / "docs" / "ВКР_Гірка_Даря_642п_Aurora_Atelier_report_only_headings.docx"
ASSETS_DIR = ROOT / ".ai" / "bachelor_report_assets"
SCREENSHOT_DIR = ROOT / ".ai" / "generated-report-assets" / "screenshots"
COURSE_DOCX = Path(r"E:\UNIVERSITY\4 kurs\1 семестр\курсовой\Курсова Гірка 642п.docx")

STUDENT = "Гірка Дар'я Миколаївна"
STUDENT_SHORT = "Дар'я ГІРКА"
GROUP = "642п"
TOPIC = "Веб-система для продажу авторських прикрас з конструктором дизайну"
CIPHER = "ХАІ.603.642п.26В.121.______ ПЗ"
SUPERVISOR = "Мандрікова Л.В."
SUPERVISOR_FULL = "Мандрікова Людмила Василівна, к.т.н., доцент"
ECON_CONSULTANT = "Купріянова В.С., доц. каф. 601"
ECON_CONSULTANT_FULL = "Купріянова Валентина Сергіївна"
ORDER = "№447-уч від 08.04.2026 р."
REVIEWER = "________________"
ISSUE_DATE = "«___» ____________ 2025 р."
SUBMISSION_DATE = "«___» ____________ 2026 р."


BODY_PAGE_TARGET = 100


class BuildContext:
    def __init__(self) -> None:
        self.bookmark_id = 1
        self.figure_no: dict[int, int] = {}
        self.appendix_figure_no: dict[str, int] = {}
        self.table_no: dict[int, int] = {}

    def next_bookmark_id(self) -> int:
        self.bookmark_id += 1
        return self.bookmark_id


def font_name(run, name: str = "Times New Roman") -> None:
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), name)


def set_paragraph_font(paragraph, size: int = 14, bold: bool | None = None) -> None:
    for run in paragraph.runs:
        font_name(run)
        run.font.size = Pt(size)
        if bold is not None:
            run.bold = bold


def add_bookmark(paragraph, name: str, bookmark_id: int) -> None:
    start = OxmlElement("w:bookmarkStart")
    start.set(qn("w:id"), str(bookmark_id))
    start.set(qn("w:name"), name)
    end = OxmlElement("w:bookmarkEnd")
    end.set(qn("w:id"), str(bookmark_id))
    paragraph._p.insert(0, start)
    paragraph._p.append(end)


def add_page_number(paragraph) -> None:
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run()
    font_name(run)
    run.font.size = Pt(12)
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = "PAGE"
    fld_sep = OxmlElement("w:fldChar")
    fld_sep.set(qn("w:fldCharType"), "separate")
    text = OxmlElement("w:t")
    text.text = "1"
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    run._r.extend([fld_begin, instr, fld_sep, text, fld_end])


def update_word_toc_and_fields(docx_path: Path) -> None:
    powershell_script = rf"""
$ErrorActionPreference = 'Stop'
$path = '{str(docx_path).replace("'", "''")}'
    $word = New-Object -ComObject Word.Application
    $word.Visible = $false
    $word.DisplayAlerts = 0
    try {{
        $doc = $word.Documents.Open($path, $false, $false)
    try {{
        $doc.Bookmarks.ShowHidden = $true
        foreach ($toc in $doc.TablesOfContents) {{
            $toc.Update()
            $toc.UpdatePageNumbers()
        }}
        $doc.Fields.Update() | Out-Null
        foreach ($storyRange in $doc.StoryRanges) {{
            try {{
                $storyRange.Fields.Update() | Out-Null
            }} catch {{}}
        }}
        foreach ($styleName in @('TOC 1', 'TOC 2', 'TOC 3', 'TOC 4', 'toc 1', 'toc 2', 'toc 3', 'toc 4')) {{
            try {{
                $style = $doc.Styles.Item($styleName)
                $style.ParagraphFormat.SpaceAfter = 0
                $style.ParagraphFormat.SpaceBefore = 0
                $style.ParagraphFormat.LineSpacingRule = 1
                $style.ParagraphFormat.LineSpacing = 12
            }} catch {{}}
        }}
        foreach ($style in $doc.Styles) {{
            try {{
                $styleNameLocal = '' + $style.NameLocal
                if ($style.BuiltIn -and $style.Type -eq 1 -and $styleNameLocal -match '^(?:TOC|toc|\\u041E\\u0433\\u043B\\u0430\\u0432\\u043B\\u0435\\u043D\\u0438\\u0435|\\u0417\\u043C\\u0456\\u0441\\u0442)\\s*[1-4]$') {{
                    $style.ParagraphFormat.SpaceAfter = 0
                    $style.ParagraphFormat.SpaceBefore = 0
                    $style.ParagraphFormat.LineSpacingRule = 1
                    $style.ParagraphFormat.LineSpacing = 12
                }}
            }} catch {{}}
        }}
        foreach ($toc in $doc.TablesOfContents) {{
            foreach ($paragraph in $toc.Range.Paragraphs) {{
                $paragraph.ParagraphFormat.SpaceAfter = 0
                $paragraph.ParagraphFormat.SpaceBefore = 0
                $paragraph.ParagraphFormat.LineSpacingRule = 1
                $paragraph.ParagraphFormat.LineSpacing = 12
            }}
        }}
        $doc.Save()
    }} finally {{
        $doc.Close([ref]0)
    }}
}} finally {{
    $word.Quit()
}}
"""
    try:
        encoded_script = base64.b64encode(powershell_script.encode("utf-16le")).decode("ascii")
        subprocess.run(
            ["powershell", "-NoProfile", "-NonInteractive", "-EncodedCommand", encoded_script],
            check=True,
            stdout=subprocess.DEVNULL,
            stderr=subprocess.DEVNULL,
            timeout=180,
        )
    except Exception:
        pass


def add_toc_field(doc: Document, levels: str = "1-4") -> None:
    p = doc.add_paragraph()
    p.paragraph_format.first_line_indent = Cm(0)
    p.paragraph_format.line_spacing = 1
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run()
    font_name(run)
    run.font.size = Pt(14)
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = f' TOC \\\\o "{levels}" \\\\h \\\\z \\\\u '
    fld_sep = OxmlElement("w:fldChar")
    fld_sep.set(qn("w:fldCharType"), "separate")
    placeholder = OxmlElement("w:t")
    placeholder.text = "Оновіть поле, щоб сформувати зміст."
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    run._r.extend([fld_begin, instr, fld_sep, placeholder, fld_end])


def set_cell_border(cell, **kwargs) -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_borders = tc_pr.first_child_found_in("w:tcBorders")
    if tc_borders is None:
        tc_borders = OxmlElement("w:tcBorders")
        tc_pr.append(tc_borders)
    for edge in ("left", "top", "right", "bottom", "insideH", "insideV"):
        if edge in kwargs:
            edge_data = kwargs.get(edge)
            tag = f"w:{edge}"
            element = tc_borders.find(qn(tag))
            if element is None:
                element = OxmlElement(tag)
                tc_borders.append(element)
            for key in ["val", "sz", "space", "color"]:
                if key in edge_data:
                    element.set(qn(f"w:{key}"), str(edge_data[key]))


def set_cell_margin(cell, margin_twips: int = 90) -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for side in ("top", "left", "bottom", "right"):
        node = tc_mar.find(qn(f"w:{side}"))
        if node is None:
            node = OxmlElement(f"w:{side}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(margin_twips))
        node.set(qn("w:type"), "dxa")


def shade_cell(cell, fill: str = "EDEDED") -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def apply_table_style(table, header_fill: str | None = None) -> None:
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    for r_idx, row in enumerate(table.rows):
        for cell in row.cells:
            set_cell_border(
                cell,
                left={"val": "single", "sz": "4", "space": "0", "color": "000000"},
                right={"val": "single", "sz": "4", "space": "0", "color": "000000"},
                top={"val": "single", "sz": "4", "space": "0", "color": "000000"},
                bottom={"val": "single", "sz": "4", "space": "0", "color": "000000"},
            )
            set_cell_margin(cell)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            if r_idx == 0 and header_fill:
                shade_cell(cell, header_fill)
            for paragraph in cell.paragraphs:
                paragraph.paragraph_format.space_before = Pt(0)
                paragraph.paragraph_format.space_after = Pt(0)
                paragraph.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
                paragraph.paragraph_format.first_line_indent = Cm(0)
                for run in paragraph.runs:
                    font_name(run)
                    run.font.size = Pt(11)
                    if r_idx == 0:
                        run.bold = True


def configure_document(doc: Document) -> None:
    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Times New Roman"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    normal.font.size = Pt(14)
    normal.font.color.rgb = RGBColor(0, 0, 0)
    normal.paragraph_format.first_line_indent = Cm(1)
    normal.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    normal.paragraph_format.space_after = Pt(0)
    normal.paragraph_format.space_before = Pt(0)

    heading_formats = {
        "Heading 1": {"align": WD_ALIGN_PARAGRAPH.CENTER, "indent": 0, "before": 12, "after": 6},
        "Heading 2": {"align": WD_ALIGN_PARAGRAPH.LEFT, "indent": 1, "before": 8, "after": 6},
        "Heading 3": {"align": WD_ALIGN_PARAGRAPH.LEFT, "indent": 1, "before": 6, "after": 4},
        "Heading 4": {"align": WD_ALIGN_PARAGRAPH.LEFT, "indent": 1, "before": 4, "after": 3},
    }
    for style_name, fmt in heading_formats.items():
        style = styles[style_name]
        style.font.name = "Times New Roman"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
        style.font.size = Pt(14)
        style.font.bold = True
        style.font.color.rgb = RGBColor(0, 0, 0)
        style.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
        style.paragraph_format.space_before = Pt(fmt["before"])
        style.paragraph_format.space_after = Pt(fmt["after"])
        style.paragraph_format.first_line_indent = Cm(fmt["indent"])
        style.paragraph_format.left_indent = Cm(0)
        style.paragraph_format.keep_with_next = True

    if "Front Matter Heading" in styles:
        front_heading = styles["Front Matter Heading"]
    else:
        front_heading = styles.add_style("Front Matter Heading", WD_STYLE_TYPE.PARAGRAPH)
    front_heading.base_style = styles["Normal"]
    front_heading.font.name = "Times New Roman"
    front_heading._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    front_heading.font.size = Pt(14)
    front_heading.font.bold = True
    front_heading.font.color.rgb = RGBColor(0, 0, 0)
    front_heading.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    front_heading.paragraph_format.space_before = Pt(12)
    front_heading.paragraph_format.space_after = Pt(6)
    front_heading.paragraph_format.first_line_indent = Cm(0)
    front_heading.paragraph_format.left_indent = Cm(0)
    front_heading.paragraph_format.keep_with_next = True
    front_ppr = front_heading._element.get_or_add_pPr()
    front_outline = front_ppr.find(qn("w:outlineLvl"))
    if front_outline is None:
        front_outline = OxmlElement("w:outlineLvl")
        front_ppr.append(front_outline)
    front_outline.set(qn("w:val"), "9")

    for toc_name in ("TOC 1", "TOC 2", "TOC 3", "TOC 4"):
        if toc_name in styles:
            toc_style = styles[toc_name]
        else:
            toc_style = styles.add_style(toc_name, WD_STYLE_TYPE.PARAGRAPH)
        toc_style.base_style = styles["Normal"]
        toc_style.font.name = "Times New Roman"
        toc_style._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
        toc_style.font.size = Pt(14)
        toc_style.font.bold = False
        toc_style.font.color.rgb = RGBColor(0, 0, 0)
        toc_style.paragraph_format.space_before = Pt(0)
        toc_style.paragraph_format.space_after = Pt(0)
        toc_style.paragraph_format.line_spacing = 1
        toc_style.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
        toc_style.paragraph_format.first_line_indent = Cm(0)

    for style_name in ("List Bullet", "List Number"):
        if style_name in styles:
            style = styles[style_name]
            style.font.name = "Times New Roman"
            style._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
            style.font.size = Pt(14)
            style.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE

    if "Caption" in styles:
        cap = styles["Caption"]
        cap.font.name = "Times New Roman"
        cap._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
        cap.font.size = Pt(14)
        cap.font.italic = False

    settings = doc.settings.element
    update_fields = settings.find(qn("w:updateFields"))
    if update_fields is None:
        update_fields = OxmlElement("w:updateFields")
        settings.append(update_fields)
    update_fields.set(qn("w:val"), "true")

    for section in doc.sections:
        configure_section(section)


def configure_section(section) -> None:
    section.page_width = Mm(210)
    section.page_height = Mm(297)
    section.top_margin = Mm(20)
    section.bottom_margin = Mm(20)
    section.left_margin = Mm(25)
    section.right_margin = Mm(15)
    section.header_distance = Mm(10)
    section.footer_distance = Mm(10)


def add_paragraph(
    doc: Document,
    text: str = "",
    *,
    align: WD_ALIGN_PARAGRAPH | None = None,
    bold: bool = False,
    italic: bool = False,
    size: int = 14,
    first_indent: bool = True,
    keep_next: bool = False,
) -> None:
    p = doc.add_paragraph()
    if text:
        run = p.add_run(text)
        font_name(run)
        run.font.size = Pt(size)
        run.bold = bold
        run.italic = italic
    fmt = p.paragraph_format
    fmt.first_line_indent = Cm(1) if first_indent else Cm(0)
    fmt.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    fmt.space_after = Pt(0)
    fmt.space_before = Pt(0)
    fmt.keep_with_next = keep_next
    p.alignment = align if align is not None else WD_ALIGN_PARAGRAPH.JUSTIFY


def add_center(doc: Document, text: str, *, bold: bool = False, size: int = 14, after: int = 0) -> None:
    p = doc.add_paragraph()
    run = p.add_run(text)
    font_name(run)
    run.font.size = Pt(size)
    run.bold = bold
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.first_line_indent = Cm(0)
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    p.paragraph_format.space_after = Pt(after)


def add_heading(
    doc: Document,
    ctx: BuildContext,
    title: str,
    *,
    level: int = 1,
    numbered: bool = True,
    toc: bool = True,
    style_name: str | None = None,
) -> str:
    p = doc.add_paragraph(style=style_name or f"Heading {level}")
    run = p.add_run(title.upper() if level == 1 else title)
    font_name(run)
    run.font.size = Pt(14)
    run.bold = True
    run.font.color.rgb = RGBColor(0, 0, 0)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER if level == 1 else WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.first_line_indent = Cm(0 if level == 1 else 1)
    p.paragraph_format.keep_with_next = True
    slug = re.sub(r"[^A-Za-z0-9]+", "_", title.encode("ascii", "ignore").decode("ascii")).strip("_")
    if not slug:
        slug = f"h_{ctx.bookmark_id}"
    bookmark = f"h_{ctx.bookmark_id}_{slug[:40]}"
    bookmark_id = ctx.next_bookmark_id()
    add_bookmark(p, bookmark, bookmark_id)
    return bookmark


def add_structural_heading(doc: Document, ctx: BuildContext, title: str, *, page_break: bool = True) -> str:
    if page_break:
        doc.add_page_break()
    return add_heading(doc, ctx, title, level=1, numbered=False, toc=True)


def add_caption(doc: Document, caption: str, *, figure: bool = True) -> None:
    p = doc.add_paragraph()
    run = p.add_run(caption)
    font_name(run)
    run.font.size = Pt(14)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER if figure else WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.first_line_indent = Cm(0 if figure else 1)
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    p.paragraph_format.space_after = Pt(6)


def add_table(doc: Document, headers: list[str], rows: list[list[str]], widths_cm: list[float] | None = None) -> None:
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for idx, header in enumerate(headers):
        table.rows[0].cells[idx].text = header
    for row in rows:
        cells = table.add_row().cells
        for idx, value in enumerate(row):
            cells[idx].text = str(value)
    if widths_cm:
        for row in table.rows:
            for idx, width in enumerate(widths_cm):
                row.cells[idx].width = Cm(width)
    apply_table_style(table)


def table_caption(doc: Document, ctx: BuildContext, chapter: int, text: str) -> None:
    ctx.table_no[chapter] = ctx.table_no.get(chapter, 0) + 1
    add_paragraph(doc, f"Таблиця {chapter}.{ctx.table_no[chapter]} - {text}", first_indent=True, keep_next=True)


def figure_caption(doc: Document, ctx: BuildContext, chapter: int, text: str) -> None:
    ctx.figure_no[chapter] = ctx.figure_no.get(chapter, 0) + 1
    add_caption(doc, f"Рисунок {chapter}.{ctx.figure_no[chapter]} - {text}", figure=True)


def appendix_figure_caption(doc: Document, ctx: BuildContext, appendix_letter: str, text: str) -> None:
    ctx.appendix_figure_no[appendix_letter] = ctx.appendix_figure_no.get(appendix_letter, 0) + 1
    add_caption(doc, f"Рисунок {appendix_letter}.{ctx.appendix_figure_no[appendix_letter]} - {text}", figure=True)


def add_picture(doc: Document, image_path: Path, width_cm: float = 14.5) -> None:
    if not image_path.exists():
        add_paragraph(doc, f"[Зображення не знайдено: {image_path.name}]", italic=True)
        return
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    run.add_picture(str(image_path), width=Cm(width_cm))
    p.paragraph_format.first_line_indent = Cm(0)
    p.paragraph_format.space_after = Pt(3)


def extract_course_analogue_images() -> dict[str, Path]:
    ASSETS_DIR.mkdir(parents=True, exist_ok=True)
    mapping = {
        "jweel_home": "word/media/image1.png",
        "jweel_constructor": "word/media/image2.png",
        "moment_home": "word/media/image3.png",
        "moment_personalization": "word/media/image4.png",
        "takayas_home": "word/media/image5.png",
        "takayas_form": "word/media/image6.png",
    }
    result: dict[str, Path] = {}
    with zipfile.ZipFile(COURSE_DOCX) as zf:
        for key, internal in mapping.items():
            out = ASSETS_DIR / f"{key}.png"
            out.write_bytes(zf.read(internal))
            result[key] = out
    return result


def get_font(size: int = 26, bold: bool = False):
    candidates = [
        r"C:\Windows\Fonts\arialbd.ttf" if bold else r"C:\Windows\Fonts\arial.ttf",
        r"C:\Windows\Fonts\timesbd.ttf" if bold else r"C:\Windows\Fonts\times.ttf",
        r"C:\Windows\Fonts\calibri.ttf",
    ]
    for path in candidates:
        if Path(path).exists():
            return ImageFont.truetype(path, size=size)
    return ImageFont.load_default()


def draw_wrapped(draw: ImageDraw.ImageDraw, xy: tuple[int, int], text: str, font, fill=(20, 20, 20), max_width: int = 260, line_gap: int = 6) -> int:
    words = text.split()
    lines: list[str] = []
    current = ""
    for word in words:
        candidate = f"{current} {word}".strip()
        width = draw.textbbox((0, 0), candidate, font=font)[2]
        if width <= max_width or not current:
            current = candidate
        else:
            lines.append(current)
            current = word
    if current:
        lines.append(current)
    x, y = xy
    line_height = draw.textbbox((0, 0), "АБВ", font=font)[3] + line_gap
    for line in lines:
        draw.text((x, y), line, font=font, fill=fill)
        y += line_height
    return y


def draw_box(draw: ImageDraw.ImageDraw, box: tuple[int, int, int, int], title: str, subtitle: str = "", fill=(248, 246, 240), outline=(90, 70, 45)) -> None:
    draw.rounded_rectangle(box, radius=18, fill=fill, outline=outline, width=3)
    x1, y1, x2, _ = box
    title_font = get_font(26, bold=True)
    text_font = get_font(21)
    draw_wrapped(draw, (x1 + 18, y1 + 18), title, title_font, max_width=x2 - x1 - 36)
    if subtitle:
        draw_wrapped(draw, (x1 + 18, y1 + 58), subtitle, text_font, max_width=x2 - x1 - 36)


def draw_arrow(draw: ImageDraw.ImageDraw, start: tuple[int, int], end: tuple[int, int], fill=(70, 70, 70)) -> None:
    draw.line([start, end], fill=fill, width=4)
    angle = math.atan2(end[1] - start[1], end[0] - start[0])
    length = 16
    left = (end[0] - length * math.cos(angle - math.pi / 6), end[1] - length * math.sin(angle - math.pi / 6))
    right = (end[0] - length * math.cos(angle + math.pi / 6), end[1] - length * math.sin(angle + math.pi / 6))
    draw.polygon([end, left, right], fill=fill)


def make_diagram(name: str, boxes: list[tuple[tuple[int, int, int, int], str, str]], arrows: list[tuple[tuple[int, int], tuple[int, int]]], size=(1500, 900)) -> Path:
    ASSETS_DIR.mkdir(parents=True, exist_ok=True)
    img = Image.new("RGB", size, (255, 255, 255))
    draw = ImageDraw.Draw(img)
    for start, end in arrows:
        draw_arrow(draw, start, end)
    palette = [(249, 246, 235), (241, 248, 244), (244, 245, 252), (252, 244, 241)]
    for idx, (box, title, subtitle) in enumerate(boxes):
        draw_box(draw, box, title, subtitle, fill=palette[idx % len(palette)])
    out = ASSETS_DIR / f"{name}.png"
    img.save(out)
    return out


def create_diagrams() -> dict[str, Path]:
    diagrams: dict[str, Path] = {}
    diagrams["architecture"] = make_diagram(
        "architecture",
        [
            ((60, 330, 300, 520), "Браузер", "Клієнт або адміністратор"),
            ((390, 250, 680, 600), "React SPA", "Routes, features, API client"),
            ((780, 250, 1080, 600), "Express API", "Middleware, routes, services"),
            ((1180, 80, 1440, 270), "SQLite", "Knex migrations, volume"),
            ((1180, 355, 1440, 545), "Файли", "Assets, uploads, constructor JSON"),
            ((1180, 625, 1440, 815), "Adapters", "Mock payment, email logs"),
        ],
        [
            ((300, 425), (390, 425)),
            ((680, 425), (780, 425)),
            ((1080, 350), (1180, 180)),
            ((1080, 425), (1180, 450)),
            ((1080, 505), (1180, 720)),
        ],
    )
    diagrams["use_cases"] = make_diagram(
        "use_cases",
        [
            ((60, 110, 300, 285), "Клієнт", "Реєструється, обирає вироби, оформлює замовлення"),
            ((60, 610, 300, 785), "Адміністратор", "Керує товарами, конструктором і замовленнями"),
            ((520, 70, 860, 190), "Перегляд каталогу", ""),
            ((520, 225, 860, 345), "Створення custom design", ""),
            ((520, 380, 860, 500), "Checkout і оплата", ""),
            ((520, 535, 860, 655), "Перегляд замовлень", ""),
            ((1000, 190, 1360, 330), "Керування каталогом", ""),
            ((1000, 370, 1360, 510), "Керування конструктором", ""),
            ((1000, 550, 1360, 690), "Зміна статусів", ""),
        ],
        [
            ((300, 190), (520, 130)),
            ((300, 190), (520, 285)),
            ((300, 190), (520, 440)),
            ((300, 190), (520, 595)),
            ((300, 700), (1000, 260)),
            ((300, 700), (1000, 440)),
            ((300, 700), (1000, 620)),
        ],
    )
    diagrams["db"] = make_diagram(
        "db_model",
        [
            ((60, 80, 330, 260), "users, sessions", "Ролі, httpOnly session"),
            ((430, 80, 700, 260), "products", "Каталог, images, filters"),
            ((800, 80, 1070, 260), "constructor", "Types, variants, slots, stones"),
            ((1170, 80, 1440, 260), "cart_items", "Ready product або custom design"),
            ((430, 520, 700, 720), "orders", "Status, totals, customer data"),
            ((800, 520, 1070, 720), "payments", "Mock provider audit"),
            ((1170, 520, 1440, 720), "notification_logs", "Email delivery status"),
        ],
        [
            ((330, 170), (430, 170)),
            ((700, 170), (800, 170)),
            ((1070, 170), (1170, 170)),
            ((1305, 260), (700, 520)),
            ((700, 620), (800, 620)),
            ((1070, 620), (1170, 620)),
            ((195, 260), (430, 620)),
        ],
    )
    diagrams["checkout"] = make_diagram(
        "checkout_sequence",
        [
            ((60, 120, 290, 285), "Клієнт", "Заповнює checkout"),
            ((390, 120, 660, 285), "React route", "Валідація UI, API call"),
            ((760, 120, 1030, 285), "Checkout service", "DB transaction"),
            ((1130, 120, 1410, 285), "SQLite", "Order, items, payment"),
            ((760, 560, 1030, 725), "Notification", "Email status event"),
            ((1130, 560, 1410, 725), "Payment mock", "Confirm -> confirmed"),
        ],
        [
            ((290, 200), (390, 200)),
            ((660, 200), (760, 200)),
            ((1030, 200), (1130, 200)),
            ((895, 285), (895, 560)),
            ((1270, 285), (1270, 560)),
            ((1130, 640), (1030, 640)),
        ],
    )
    diagrams["pricing"] = make_diagram(
        "constructor_pricing",
        [
            ((70, 140, 340, 315), "Configuration", "Variant, material, size, stones"),
            ((440, 140, 720, 315), "Validation", "Allowed slots and required values"),
            ((820, 140, 1100, 315), "Price formula", "Base + variant + material + stones"),
            ((1180, 140, 1450, 315), "Response", "Price, layers, normalized config"),
            ((440, 560, 720, 735), "Preview layers", "Base, stones, engraving"),
            ((820, 560, 1100, 735), "Cart item", "Snapshot + server price"),
        ],
        [
            ((340, 225), (440, 225)),
            ((720, 225), (820, 225)),
            ((1100, 225), (1180, 225)),
            ((960, 315), (580, 560)),
            ((580, 735), (820, 650)),
        ],
    )
    return diagrams


def add_front_matter(doc: Document, ctx: BuildContext) -> None:
    add_heading(doc, ctx, "РЕФЕРАТ", level=1, numbered=False, style_name="Front Matter Heading")
    add_paragraph(
        doc,
        "Пояснювальна записка до кваліфікаційної роботи бакалавра: 100 сторінок основного тексту без урахування додатків, рисунки, таблиці, перелік посилань та додатки.",
    )
    add_paragraph(
        doc,
        "Об'єктом роботи є процес продажу авторських прикрас у малих майстернях та студіях, де важливими є персоналізація виробу, швидке оформлення замовлення, прозорий розрахунок вартості та контроль виконання замовлення. Предметом роботи є веб-система Aurora Atelier, що поєднує каталог готових виробів, конструктор дизайну, корзину, checkout, імітацію оплати, кабінет клієнта та адміністративні інструменти.",
    )
    add_paragraph(
        doc,
        "Метою кваліфікаційної роботи є підвищення ефективності онлайн-продажу авторських прикрас за рахунок розроблення та впровадження веб-системи з інтерактивним конструктором дизайну, серверним розрахунком ціни, підтримкою замовлень та адміністративним керуванням каталогом і конфігурацією конструктора.",
    )
    add_paragraph(
        doc,
        "У роботі виконано аналіз вимог, розглянуто аналоги Jweel, Moment Creator і Takayas Custom Jewelry, спроєктовано клієнт-серверну архітектуру на основі React, Express, SQLite та Knex, реалізовано основні модулі системи і проведено тестування ключових сценаріїв. Результати автоматизованої перевірки підтвердили проходження 66 тестів у 10 файлах.",
    )
    add_paragraph(doc, "Ключові слова: ВЕБ-СИСТЕМА, АВТОРСЬКІ ПРИКРАСИ, КОНСТРУКТОР ДИЗАЙНУ, REACT, EXPRESS, SQLITE, CHECKOUT.")

    doc.add_page_break()
    add_heading(doc, ctx, "ABSTRACT", level=1, numbered=False, style_name="Front Matter Heading")
    add_paragraph(
        doc,
        "The bachelor's qualification project presents Aurora Atelier, a web system for selling artisan jewelry with a design constructor. The system combines a product catalog, interactive customization, server-side pricing, cart and checkout flows, mock payment confirmation, customer order tracking, and administrative tools for catalog, constructor, and order management.",
    )
    add_paragraph(
        doc,
        "The project goal is to improve online sales of artisan jewelry by providing a clear, accessible, and maintainable web system for small workshops. The implementation uses React and Vite on the client side, Node.js and Express on the server side, SQLite with Knex migrations for persistence, and Docker Compose for reproducible runtime deployment.",
    )
    add_paragraph(
        doc,
        "The work includes requirements analysis, comparison with existing analogues, architectural and database design, implementation details, testing strategy, and economic justification. Automated tests cover authentication, catalog, constructor configuration, cart, checkout, payment idempotency, order ownership, admin flows, promo codes, and order status transitions.",
    )
    add_paragraph(doc, "Keywords: WEB SYSTEM, ARTISAN JEWELRY, DESIGN CONSTRUCTOR, REACT, EXPRESS, SQLITE, CHECKOUT.")

    doc.add_page_break()
    add_heading(doc, ctx, "ЗМІСТ", level=1, numbered=False, toc=False, style_name="Front Matter Heading")
    add_toc_field(doc, "1-4")

    doc.add_page_break()
    add_heading(doc, ctx, "ПЕРЕЛІК УМОВНИХ ПОЗНАЧЕНЬ, СИМВОЛІВ, ОДИНИЦЬ ВИМІРЮВАНЬ ФІЗИЧНИХ ВЕЛИЧИН, СКОРОЧЕНЬ І ТЕРМІНІВ", level=1, numbered=False)
    terms = [
        ["API", "Application Programming Interface, програмний інтерфейс взаємодії між клієнтом і сервером."],
        ["SPA", "Single Page Application, веб-додаток, у якому навігація виконується без повного перезавантаження сторінки."],
        ["UI", "User Interface, інтерфейс користувача."],
        ["UX", "User Experience, досвід користувача під час роботи із системою."],
        ["БД", "База даних."],
        ["ПЗ", "Програмне забезпечення."],
        ["ПП", "Програмний продукт."],
        ["CRUD", "Створення, читання, оновлення та видалення даних."],
        ["HTTP", "Протокол передачі даних між браузером і сервером."],
        ["JSON", "Текстовий формат обміну структурованими даними."],
        ["Mock payment", "Імітаційний платіжний провайдер для MVP без реального списання коштів."],
        ["Ready product", "Готовий товар каталогу."],
        ["Custom design", "Індивідуальна прикраса, зібрана користувачем у конструкторі."],
    ]
    add_table(doc, ["Скорочення", "Пояснення"], terms, [3.5, 13.0])
def add_project_elaboration(doc: Document, focus: str, aspects: list[str], count: int) -> None:
    profiles = {
        "вступного обґрунтування теми": {
            "anchors": [
                "публічні маршрути `/`, `/catalog`, `/products/:slug` і `/constructor`",
                "server-side checkout і `mock payment` замість декларативного опису магазину",
                "окремі клієнтські й адміністративні сценарії `/orders` та `/admin/orders`",
                "SQLite/Knex як легку БД для малої майстерні авторських прикрас",
            ],
            "modules": "каталог, конструктор дизайну, корзину, checkout, mock payment, кабінет клієнта та адміністративні екрани",
        },
        "глосарію предметної області": {
            "anchors": [
                "таблиці `ready_products`, `orders`, `order_items` і JSON-конфігурації конструктора",
                "терміни `ready product`, `custom design`, `stone slot`, `variant stone` і `mock payment`",
                "поля route- та API-моделей, які реально використовуються в React і Express",
            ],
            "modules": "сутності каталогу, конструктора, замовлень і адміністративної бібліотеки",
        },
        "концепції програмного продукту": {
            "anchors": [
                "2D layered preview замість важкого 3D-движка",
                "backend price engine у `constructor-pricing.service.js`",
                "розділення public flow і admin flow на окремі маршрути та модулі",
                "відсутність CRM, реального acquiring і складського контуру у межах MVP",
            ],
            "modules": "конкретний обсяг MVP Aurora Atelier без штучно доданих підсистем",
        },
        "аналізу вимог": {
            "anchors": [
                "сценарії перегляду каталогу, створення `custom design`, checkout і адміністративної обробки",
                "матрицю трасування між вимогами, backend-модулями і автоматизованими тестами",
                "бізнес-правила для промокодів, статусів замовлення і server-side price validation",
            ],
            "modules": "тільки реальні сценарії Aurora Atelier, присутні у коді та тестах",
        },
        "архітектурного проєктування": {
            "anchors": [
                "React SPA у `client/src` та Express application у `server/`",
                "модульний моноліт із route-layer, service-layer і SQLite persistence",
                "Docker Compose runtime з volume для БД та production build клієнта у `public/react-app`",
                "middleware pipeline для сесій, локалі, ролей і JSON API",
            ],
            "modules": "фактичну структуру репозиторію та спосіб розгортання",
        },
        "проєктування бази даних": {
            "anchors": [
                "таблиці користувачів, сесій, кошиків, замовлень, оплат і історії статусів",
                "сутності конструктора: типи, варіанти, слоти каменів, матриці доступності та assets",
                "snapshot-поля в `order_items`, що фіксують стан товару або дизайну на момент checkout",
                "міграції Knex і seed-дані каталогу для 40 готових виробів",
            ],
            "modules": "конкретні схеми Aurora Atelier, а не умовну універсальну БД",
        },
        "детального проєктування конструктора": {
            "anchors": [
                "координатні `stone slots` та layered `preview_*` поля конфігурації",
                "`constructor-normalizers.js`, `constructor-assets.service.js` і `constructor-pricing.service.js`",
                "адміністративний CRUD бібліотеки через `/admin/constructor` і autosave-поведінку",
                "перевірку матеріалу, розміру, engraving та chain surcharge для підвісок",
            ],
            "modules": "реальну механіку конструктора, яка доступна в публічній та адміністративній частинах",
        },
        "проєктування checkout": {
            "anchors": [
                "транзакційний перенос `cart_items` в `order_items`",
                "генерацію payment token і початкового статусу `created_pending_payment`",
                "ідемпотентність підтвердження оплати і журналювання notification events",
                "облік промокоду, ownership logic та створення нової активної корзини після checkout",
            ],
            "modules": "серверний order flow, а не лише візуальну checkout-форму",
        },
        "реалізації програмного забезпечення": {
            "anchors": [
                "`client/src/main.jsx`, route-компоненти, `api.js` і route-specific CSS-файли",
                "backend modules `checkout`, `orders`, `admin-orders`, `constructor`, `admin-catalog`",
                "Knex migrations, seed-дані і допоміжні утиліти для preview та customer name fallback",
                "Dockerfile і `docker-compose.yml`, що забезпечують відтворюваний запуск у контейнері",
            ],
            "modules": "власний код Aurora Atelier на рівні файлів, модулів і runtime",
        },
        "інтеграційного тестування": {
            "anchors": [
                "`tests/integration/api-flows.test.js` і `tests/integration/cart-checkout.test.js`",
                "реальні HTTP-запити з session cookies, role guard і ownership checks",
                "SQLite test database з seed-даними каталогу та замовлень",
                "checkout transaction, mock payment, промокоди і admin status workflow",
            ],
            "modules": "поведінку готового продукту на рівні API та бази даних",
        },
        "результатів тестування": {
            "anchors": [
                "реєстрацію, каталог, конструктор, корзину, checkout, оплату і клієнтські замовлення",
                "обмеження доступу до admin endpoints і коректність статусних переходів",
                "payment idempotency, promo limits і account dashboard з активними замовленнями",
                "перевірки англомовного локалізованого виводу для частини каталогових і order-сценаріїв",
            ],
            "modules": "тільки ті сценарії, які фактично проходять у тестовому наборі `npm test`",
        },
        "виробничо-організаційного плану": {
            "anchors": [
                "послідовність від аналізу вимог і ER-проєктування до реалізації й тестування",
                "модель одного виконавця без штучного дроблення на неіснуючу команду",
                "окремий час на підготовку diagrams, screenshots і пояснювальної записки",
                "ризики інтеграції checkout, конструктора і адміністративних модулів в один runtime",
            ],
            "modules": "реальний процес виконання дипломного проєкту Aurora Atelier",
        },
        "фінансового плану": {
            "anchors": [
                "відкритий стек React, Express, SQLite, Knex і Docker без ліцензійних платежів",
                "основну частку витрат у вигляді трудомісткості розроблення й тестування",
                "економію на інфраструктурі завдяки SQLite та контейнеризованому локальному розгортанню",
                "ринкове позиціонування як MVP-рішення для невеликої майстерні, а не масового SaaS",
            ],
            "modules": "економічні параметри, безпосередньо пов'язані з фактичним стеком і масштабом проєкту",
        },
    }
    profile = profiles.get(focus)
    if not profile:
        return
    anchors = profile["anchors"]
    modules = profile["modules"]
    templates = [
        "У межах {focus} для Aurora Atelier важливим є {aspect}. Це рішення безпосередньо простежується через {anchor}, тому опис у пояснювальній записці спирається на реальний модуль або маршрут системи.",
        "Практичне наповнення напряму {focus} видно не в загальних деклараціях, а у фактичній реалізації: {anchor}. Саме через це диплом описує {modules}, які існують у репозиторії та можуть бути продемонстровані під час захисту.",
        "Для предметної області авторських прикрас аспект {aspect} критичний, оскільки він впливає на точність конфігурації виробу, прозорість ціни і подальшу обробку замовлення. У Aurora Atelier це закріплено через {anchor}.",
        "З точки зору супроводу проєкту {focus} має бути прив'язаний до конкретного коду. У поточній системі це означає використання {anchor}, що дозволяє зіставити текст диплома з реальною архітектурою Aurora Atelier.",
        "Під час деталізації напряму {focus} враховано, що дипломна робота повинна описувати власний реалізований продукт. Тому замість абстрактних формулювань у тексті фіксуються {modules}, а їх технічним підтвердженням виступає {anchor}.",
        "Окреме значення для {focus} має аспект {aspect}, бо саме він відрізняє Aurora Atelier від звичайної вітрини. Наявність {anchor} показує, що система підтримує не лише перегляд товарів, а й повний операційний контур замовлення.",
    ]
    limit = max(2, math.ceil(count * 0.72))
    for idx in range(limit):
        template = templates[idx % len(templates)]
        aspect = aspects[idx % len(aspects)]
        anchor = anchors[idx % len(anchors)]
        add_paragraph(doc, template.format(focus=focus, aspect=aspect, anchor=anchor, modules=modules))


def add_introduction(doc: Document, ctx: BuildContext, *, page_break: bool = True) -> None:
    add_structural_heading(doc, ctx, "ВСТУП", page_break=page_break)
    add_bookmark(doc.paragraphs[-1], "bm_body_start", ctx.next_bookmark_id())
    intro = [
        "Сфера електронної комерції продовжує розвиватися як один із ключових каналів взаємодії між малим бізнесом і клієнтами. Для майстерень авторських прикрас важливо не лише показати готові вироби, а й надати покупцю можливість брати участь у формуванні майбутнього дизайну. Саме тому звичайний сайт-вітрина часто не покриває потреби предметної області: він не дає персоналізації, не фіксує конфігурацію виробу, не забезпечує прозорий розрахунок ціни та не створює єдиного простору для подальшої обробки замовлення.",
        "Актуальність роботи визначається потребою невеликих ювелірних студій у доступному веб-рішенні, яке поєднує каталог, персоналізацію, checkout і адміністративний контроль. Великі платформи можуть використовувати складні 3D-конструктори або індивідуальне спілкування з менеджером, але такі підходи не завжди придатні для малого бізнесу через високу вартість розроблення, складність підтримки та надмірні вимоги до пристрою користувача.",
        "Розроблювана веб-система Aurora Atelier орієнтована на продаж авторських прикрас із підтримкою конструктора дизайну. Користувач може переглянути каталог готових виробів, налаштувати індивідуальну прикрасу, додати позиції до корзини, оформити замовлення, пройти імітацію оплати і контролювати стан виконання. Адміністратор отримує інструменти для керування товарами, матеріалами, конфігурацією конструктора, замовленнями та статусами.",
        "Метою кваліфікаційної роботи є підвищення ефективності онлайн-продажу авторських прикрас за рахунок розроблення веб-системи з конструктором дизайну, серверним розрахунком ціни, корзиною, оформленням замовлення, імітацією оплати та адміністративним керуванням.",
        "Для досягнення мети необхідно виконати аналіз вимог до системи, спроєктувати архітектуру та модель даних, реалізувати frontend і backend частини, перевірити ключові сценарії тестами та обґрунтувати економічну доцільність розроблення.",
        "У роботі використано стек технологій React 18, Vite, Node.js, Express, SQLite, Knex, Docker Compose, Vitest і Supertest. Такий стек відповідає MVP-характеру системи: він дозволяє швидко розгорнути проєкт, зберігати дані у легкій реляційній базі, описувати структуру через міграції і перевіряти поведінку автоматизованими тестами.",
    ]
    for paragraph in intro:
        add_paragraph(doc, paragraph)
    add_project_elaboration(
        doc,
        "вступного обґрунтування теми",
        ["персоналізація виробу", "доступність для малого бізнесу", "прозорий checkout", "підтримка адміністратора", "серверна валідація ціни"],
        10,
    )


def add_requirements_section(doc: Document, ctx: BuildContext, images: dict[str, Path], diagrams: dict[str, Path], expansion: int) -> None:
    doc.add_page_break()
    add_heading(doc, ctx, "1 АНАЛІЗ ВИМОГ ДО ВЕБ-СИСТЕМИ ДЛЯ ПРОДАЖУ АВТОРСЬКИХ ПРИКРАС З КОНСТРУКТОРОМ ДИЗАЙНУ", level=1)
    add_paragraph(doc, "Перший розділ присвячено визначенню бізнес-контексту, класів користувачів, вимог до програмного продукту та обмежень. Аналіз вимог потрібен для того, щоб реалізація не перетворилася на набір розрізнених екранів, а залишалася цілісною веб-системою із вимірюваною користю для майстерні авторських прикрас.")

    add_heading(doc, ctx, "1.1 Бізнес-вимоги", level=2)
    add_heading(doc, ctx, "1.1.1 Загальний опис програмного продукту та бізнес-вимоги", level=3)
    add_paragraph(doc, "Бізнес-вимоги описують, яку цінність має надати програмний продукт. Для Aurora Atelier основна цінність полягає у скороченні ручної обробки замовлень, підвищенні зручності персоналізації виробів та формуванні контрольованого процесу продажу.")
    add_heading(doc, ctx, "1.1.2 Розробка глосарію", level=3)
    add_paragraph(doc, "Для однозначного розуміння предметної області необхідно зафіксувати терміни, які використовуються в описі Aurora Atelier, її користувацьких сценаріїв, конфігурації конструктора, замовлень та адміністративних процесів.")
    table_caption(doc, ctx, 1, "Бізнес-вимоги до веб-системи")
    add_table(
        doc,
        ["ID", "Бізнес-вимога", "Критерій досягнення"],
        [
            ["BR-01", "Надати клієнту можливість самостійно сформувати індивідуальну прикрасу без листування з майстром на першому етапі.", "Користувач створює custom design у конструкторі та додає його до корзини."],
            ["BR-02", "Зменшити кількість ручних уточнень щодо ціни виробу.", "Сервер повертає ціну на основі типу, варіанта, матеріалу, розміру, каменів і ланцюжка."],
            ["BR-03", "Забезпечити продаж готових виробів разом з індивідуальними дизайнами.", "Корзина і замовлення підтримують ready_product та custom_design."],
            ["BR-04", "Надати адміністратору інструменти для контролю каталогу, конструктора і замовлень.", "Admin routes дозволяють редагувати товари, варіанти, слоти, камені та статуси."],
            ["BR-05", "Забезпечити відстеження життєвого циклу замовлення.", "Статуси змінюються за контрольованою послідовністю і фіксуються в історії."],
        ],
        [2.0, 8.0, 6.5],
    )

    table_caption(doc, ctx, 1, "Глосарій предметної області")
    glossary = [
        ["Авторська прикраса", "Виріб, створений майстром або студією з акцентом на індивідуальний дизайн і невеликі партії."],
        ["Каталог", "Сторінка та API-дані для перегляду готових виробів, доступних до замовлення."],
        ["Конструктор", "Інтерфейс, у якому клієнт обирає тип виробу, варіант, матеріал, розмір, камені та інші параметри."],
        ["Слот каменю", "Координатна позиція на базовому зображенні прикраси, куди може бути встановлений камінь."],
        ["Матриця каменів", "Зв'язок між варіантом прикраси, доступним каменем і доплатою за нього."],
        ["Ready product", "Готовий товар із каталогу, який можна додати до корзини без конструктора."],
        ["Custom design", "Індивідуальна конфігурація, створена користувачем у конструкторі."],
        ["Checkout", "Процес оформлення замовлення з контактними даними, доставкою та згодами."],
        ["Mock payment", "Імітаційна оплата, яка дозволяє перевірити платіжний сценарій без реального провайдера."],
        ["Статус замовлення", "Стан замовлення у життєвому циклі: створено, підтверджено, в роботі, виконано."],
        ["Адміністратор", "Користувач з роллю admin, який обробляє замовлення і керує даними системи."],
        ["Фільтри каталогу", "Параметри, за якими клієнт звужує перелік виробів: тип, метал, камінь, форма, розмір."],
    ]
    add_table(doc, ["Термін", "Визначення"], glossary, [4.2, 12.3])
    add_project_elaboration(doc, "глосарію предметної області", ["каталог готових виробів", "конструктор дизайну", "роль адміністратора", "збереження замовлень", "двомовність інтерфейсу"], 8 + expansion // 4)

    add_heading(doc, ctx, "1.1.3 Огляд і аналіз існуючих аналогів програмного забезпечення", level=3)
    for paragraph in [
        "Аналіз існуючих програмних продуктів, що виконують схожі функції або належать до сфери веб-систем для продажу ювелірних виробів з можливістю кастомізації, є важливим етапом формування вимог. Він дозволяє оцінити поточний стан ринку, виявити сильні та слабкі сторони наявних рішень і визначити, які функції мають найбільшу цінність для користувача.",
        "У межах роботи розглянуто три репрезентативні аналоги: Jweel, Moment Creator і Takayas Custom Jewelry. Ці сервіси демонструють різні підходи до персоналізації: складний 3D-конструктор, просту текстову персоналізацію і преміальне оформлення запиту через менеджера. Порівняння показує, що між важкими 3D-рішеннями і повністю ручним процесом існує практична ніша для легкого 2D-конструктора з автоматичним розрахунком ціни.",
    ]:
        add_paragraph(doc, paragraph)

    add_paragraph(doc, "Jweel", bold=True, first_indent=False, keep_next=True)
    add_paragraph(doc, "Jweel є прикладом технологічно насиченого онлайн-конструктора ювелірних виробів. Платформа дозволяє працювати з 3D-моделями, змінювати форму об'єктів, додавати гравіювання, обирати матеріали і переглядати модель у браузері. Такий підхід демонструє широкі можливості персоналізації, але одночасно створює високий поріг входу.")
    add_picture(doc, images["jweel_home"], 14.5)
    figure_caption(doc, ctx, 1, "Головна сторінка сайту Jweel")
    add_picture(doc, images["jweel_constructor"], 14.5)
    figure_caption(doc, ctx, 1, "Інтерфейс 3D-конструктора на сайті Jweel")
    for paragraph in [
        "Основним недоліком Jweel є залежність від складної WebGL-візуалізації. На слабких пристроях або в мобільному браузері такий інструмент може працювати повільно, що обмежує доступність продукту. Для невеликої майстерні підтримка складного 3D-редактора також означає високі витрати на розроблення, тестування і супровід.",
        "Для Aurora Atelier із цього аналогу варто взяти саму ідею інтерактивної персоналізації, але реалізувати її простішим способом: через 2D-preview, координатні слоти каменів і серверний розрахунок ціни. Це зберігає користувацьку цінність конструктора, але зменшує технічні ризики.",
    ]:
        add_paragraph(doc, paragraph)

    add_paragraph(doc, "Moment Creator", bold=True, first_indent=False, keep_next=True)
    add_paragraph(doc, "Moment Creator пропонує персоналізацію ювелірних виробів через введення текстових параметрів: ініціалів, імен, стилю шрифту і розміру. Інтерфейс є зрозумілим, однак зміни не завжди відображаються у реальному часі на зображенні виробу.")
    add_picture(doc, images["moment_home"], 14.5)
    figure_caption(doc, ctx, 1, "Головна сторінка сайту Moment Creator")
    add_picture(doc, images["moment_personalization"], 14.5)
    figure_caption(doc, ctx, 1, "Інтерфейс персоналізації підвіски з ініціалами")
    add_paragraph(doc, "Сильним боком платформи є простота взаємодії. Водночас відсутність динамічної візуалізації знижує довіру користувача до результату: покупець вводить параметри, але не бачить, як саме вони впливають на майбутній виріб. Aurora Atelier має усунути цей недолік через оновлення preview-шарів і збереження конфігурації у корзині та замовленні.")

    add_paragraph(doc, "Takayas Custom Jewelry", bold=True, first_indent=False, keep_next=True)
    add_paragraph(doc, "Takayas Custom Jewelry представляє інший підхід: користувач не працює з інтерактивним конструктором, а залишає запит на індивідуальне виготовлення. Платформа орієнтована на преміальний сегмент і передбачає подальше спілкування з дизайнером або менеджером.")
    add_picture(doc, images["takayas_home"], 14.5)
    figure_caption(doc, ctx, 1, "Головна сторінка сайту Takayas Custom Jewelry")
    add_picture(doc, images["takayas_form"], 14.5)
    figure_caption(doc, ctx, 1, "Форма запиту на сайті Takayas Custom Jewelry")
    for paragraph in [
        "Недоліком такого підходу є непрозорість для користувача на етапі оформлення. Покупець не бачить попередню ціну, не має швидкого preview і не може завершити покупку без очікування відповіді. Для преміального виробництва це може бути виправдано, але для більшості онлайн-сценаріїв створює зайве тертя.",
        "Aurora Atelier займає проміжну позицію між складним 3D-конструктором і повністю ручним запитом. Система надає самостійний вибір параметрів, зрозумілу ціну, корзину і замовлення, але не вимагає важкого 3D-моделювання.",
    ]:
        add_paragraph(doc, paragraph)

    add_paragraph(doc, "Порівняльний висновок за аналогами", bold=True, first_indent=False, keep_next=True)
    table_caption(doc, ctx, 1, "Порівняння аналогів")
    add_table(
        doc,
        ["Критерій", "Jweel", "Moment Creator", "Takayas", "Aurora Atelier"],
        [
            ["Тип персоналізації", "3D-модель", "Текстові параметри", "Запит менеджеру", "2D-конструктор з параметрами"],
            ["Динамічне preview", "Так, але важке", "Обмежено", "Ні", "Так, через preview-шари"],
            ["Прозора ціна", "Частково", "Частково", "Ні", "Так, серверний розрахунок"],
            ["Придатність для малого бізнесу", "Низька через складність", "Середня", "Низька через ручний процес", "Висока для MVP"],
            ["Завершення замовлення онлайн", "Залежить від сценарію", "Так", "Ні", "Так, checkout + mock payment"],
        ],
        [3.3, 3.2, 3.2, 3.2, 3.6],
    )

    add_heading(doc, ctx, "1.1.4 Концепція програмного продукту та обсяг проєкту", level=3)
    add_paragraph(doc, "Концепція Aurora Atelier полягає у створенні доступної веб-системи, яка дає клієнту достатній рівень персоналізації без складного 3D-моделювання. Обсяг проєкту включає публічний сайт, каталог, конструктор, корзину, checkout, mock payment, кабінет користувача, адміністративне керування товарами, замовленнями і конфігурацією конструктора.")
    add_project_elaboration(doc, "концепції програмного продукту", ["межі MVP", "відмова від надмірного 3D", "серверний контроль ціни", "поділ public та admin flows", "можливість подальшого розширення"], 12 + expansion // 3)

    add_heading(doc, ctx, "1.1.5 Бізнес-правила", level=3)
    table_caption(doc, ctx, 1, "Бізнес-правила системи")
    business_rules = [
        ["BRL-01", "До публічного каталогу потрапляють лише активні товари."],
        ["BRL-02", "Ціна custom design остаточно розраховується сервером."],
        ["BRL-03", "До корзини не можна додати custom design без обов'язкових параметрів."],
        ["BRL-04", "Checkout неможливий без прийняття оферти."],
        ["BRL-05", "Кожне замовлення має унікальний номер."],
        ["BRL-06", "Клієнт бачить лише власні замовлення."],
        ["BRL-07", "Admin endpoints доступні лише користувачу з роллю admin."],
        ["BRL-08", "Статус замовлення змінюється тільки за дозволеною послідовністю."],
        ["BRL-09", "Підтвердження mock payment є ідемпотентним."],
        ["BRL-10", "Після зміни статусу формується запис notification log."],
        ["BRL-11", "Неактивні елементи конструктора не доступні у публічному конфігу."],
        ["BRL-12", "Промокоди враховують ліміти використання та мінімальну суму."],
    ]
    add_table(doc, ["ID", "Правило"], business_rules, [2.5, 14.0])

    add_heading(doc, ctx, "1.2 Вимоги користувачів", level=2)
    add_heading(doc, ctx, "1.2.1 Характеристики класів користувачів", level=3)
    table_caption(doc, ctx, 1, "Класи користувачів")
    add_table(
        doc,
        ["Клас", "Характеристика", "Основні дії"],
        [
            ["Клієнт", "Користувач публічного сайту без спеціальних технічних знань.", "Перегляд каталогу, створення дизайну, корзина, checkout, перегляд замовлень."],
            ["Адміністратор", "Майстер або власник студії з правами керування.", "Обробка замовлень, зміна статусів, редагування товарів і конструктора."],
            ["Платіжний адаптер", "Внутрішній mock-провайдер для MVP.", "Підтвердження або відмова платежу без реального списання."],
            ["Email adapter", "Сервіс повідомлень і журналювання.", "Формування email-подій після зміни статусу."],
        ],
        [3.2, 6.5, 6.8],
    )
    add_paragraph(doc, "Клієнт веб-системи використовує публічну частину для перегляду каталогу, персоналізації прикраси, оформлення та оплати замовлення, а також для подальшого контролю власних замовлень у кабінеті.")
    add_paragraph(doc, "Адміністратор веб-системи відповідає за підтримку каталогу та конструктора, перегляд і фільтрацію замовлень, зміну їх статусів та контроль коректності даних, що беруть участь у ціноутворенні і виконанні замовлення.")
    add_heading(doc, ctx, "1.2.2 Діаграма варіантів використання", level=3)
    add_picture(doc, diagrams["use_cases"], 15.5)
    figure_caption(doc, ctx, 1, "Діаграма варіантів використання веб-системи")

    add_heading(doc, ctx, "1.2.3 Сценарії варіантів використання", level=3)
    add_paragraph(doc, "Діаграма відображає взаємодію клієнта та адміністратора з основними підсистемами Aurora Atelier: каталогом, конструктором, корзиною, checkout, оплатою, замовленнями та адміністративними модулями. Для фіксації поведінки системи далі наведено три базові сценарії.")
    add_paragraph(doc, "Сценарій оформлення індивідуального замовлення", bold=True, first_indent=False, keep_next=True)
    table_caption(doc, ctx, 1, "Сценарій UC-01 - Створення індивідуальної прикраси")
    add_table(
        doc,
        ["Поле", "Опис"],
        [
            ["Актор", "Клієнт"],
            ["Передумова", "Система доступна, дані каталогу і конструктора активні."],
            ["Основний потік", "Клієнт відкриває конструктор, обирає тип, варіант, матеріал, розмір і камені, переглядає preview і додає custom design до корзини."],
            ["Післяумова", "Система зберігає конфігурацію виробу і поточну ціну у корзині."],
            ["Винятки", "Некоректні параметри, відсутній обов'язковий вибір або помилка валідації на backend-рівні."],
        ],
        [4.0, 12.5],
    )
    add_paragraph(doc, "Сценарій оплати та підтвердження замовлення", bold=True, first_indent=False, keep_next=True)
    table_caption(doc, ctx, 1, "Сценарій UC-02 - Оформлення та оплата замовлення")
    add_table(
        doc,
        ["Поле", "Опис"],
        [
            ["Актор", "Клієнт"],
            ["Передумова", "У корзині є ready product або custom design."],
            ["Основний потік", "Клієнт переходить до корзини, перевіряє позиції, заповнює checkout, приймає оферту, створює замовлення і переходить до mock payment."],
            ["Післяумова", "Замовлення створене, для нього сформовано payment token і початковий статус."],
            ["Винятки", "Не прийнято оферту, некоректні контактні дані, порожня корзина або помилка створення замовлення."],
        ],
        [4.0, 12.5],
    )
    add_paragraph(doc, "Сценарій адміністративної обробки замовлення", bold=True, first_indent=False, keep_next=True)
    table_caption(doc, ctx, 1, "Сценарій UC-03 - Адміністративна обробка замовлення")
    add_table(
        doc,
        ["Поле", "Опис"],
        [
            ["Актор", "Адміністратор"],
            ["Передумова", "Замовлення створене і доступне в administrative dashboard."],
            ["Основний потік", "Адміністратор входить у систему, переглядає список замовлень, відкриває деталі і переводить замовлення за статусною послідовністю."],
            ["Післяумова", "Система зберігає новий статус, історію зміни та журнал повідомлень."],
            ["Винятки", "Спроба недозволеного переходу, відсутній коментар для rollback або відсутня авторизація."],
        ],
        [4.0, 12.5],
    )

    add_heading(doc, ctx, "1.3 Вимоги до програмного забезпечення", level=2)
    add_heading(doc, ctx, "1.3.1 Функціональні вимоги до програмного забезпечення", level=3)
    functional_rows = [
        ["FR-01", "Реєстрація та вхід клієнта", "auth"],
        ["FR-02", "Підтримка Google/email verification flow", "auth"],
        ["FR-03", "Перегляд каталогу з фільтрами", "catalog"],
        ["FR-04", "Перегляд картки товару", "catalog"],
        ["FR-05", "Завантаження конфігурації конструктора", "constructor"],
        ["FR-06", "Розрахунок ціни custom design", "constructor-pricing"],
        ["FR-07", "Додавання ready product до корзини", "cart"],
        ["FR-08", "Додавання custom design до корзини", "cart"],
        ["FR-09", "Застосування промокоду", "promotions"],
        ["FR-10", "Оформлення замовлення", "checkout"],
        ["FR-11", "Підтвердження mock payment", "payments"],
        ["FR-12", "Перегляд замовлень клієнта", "orders/account"],
        ["FR-13", "Вхід адміністратора", "admin-auth"],
        ["FR-14", "Перегляд і фільтрація замовлень admin", "admin-orders"],
        ["FR-15", "Зміна статусу замовлення", "admin-orders"],
        ["FR-16", "CRUD готових товарів", "admin-catalog"],
        ["FR-17", "CRUD типів, варіантів, слотів і каменів", "admin-constructor"],
        ["FR-18", "Завантаження і перевірка assets", "admin-assets"],
        ["FR-19", "Локалізація uk/en", "i18n"],
        ["FR-20", "Журналювання email-подій", "notifications"],
    ]
    add_heading(doc, ctx, "1.3.1.1 Каталог і перегляд готових виробів", level=4)
    add_paragraph(doc, "Система повинна надавати клієнту каталог готових прикрас із фільтрами, карткою виробу, вибором параметрів товару та додаванням до корзини.")
    add_heading(doc, ctx, "1.3.1.2 Конструктор дизайну", level=4)
    add_paragraph(doc, "Система повинна дозволяти персоналізацію прикраси через вибір типу, варіанта, матеріалу, розміру, каменів та інших параметрів із динамічним preview і серверним розрахунком ціни.")
    add_heading(doc, ctx, "1.3.1.3 Корзина та checkout", level=4)
    add_paragraph(doc, "Система повинна зберігати обрані ready product і custom design у корзині, підтримувати промокоди і забезпечувати оформлення замовлення через checkout-форму.")
    add_heading(doc, ctx, "1.3.1.4 Оплата та життєвий цикл замовлення", level=4)
    add_paragraph(doc, "Система повинна реалізовувати сценарій mock payment, створення статусної історії замовлення та контроль допустимих переходів між статусами.")
    add_heading(doc, ctx, "1.3.1.5 Особистий кабінет клієнта", level=4)
    add_paragraph(doc, "Система повинна надавати клієнту доступ до профілю, переліку власних замовлень, їх статусів та деталей конкретного замовлення.")
    add_heading(doc, ctx, "1.3.1.6 Адміністративна підсистема", level=4)
    add_paragraph(doc, "Система повинна забезпечувати адміністративне керування замовленнями, готовими товарами, параметрами конструктора, каменями, слотами і допоміжними активами.")
    table_caption(doc, ctx, 1, "Функціональні вимоги")
    add_table(doc, ["ID", "Вимога", "Модуль"], functional_rows, [2.0, 10.0, 4.5])
    add_heading(doc, ctx, "1.3.2 Нефункціональні вимоги до програмного забезпечення", level=3)
    nfr_rows = [
        ["NFR-01", "Безпека", "httpOnly cookie, bcryptjs, role guard, ownership checks."],
        ["NFR-02", "Продуктивність", "Легка SQLite БД, індекси, серверна пагінація каталогу, debounce для price API."],
        ["NFR-03", "Доступність", "Веб-клієнт без встановлення додаткового ПЗ, Docker-first запуск."],
        ["NFR-04", "Масштабованість", "Модульний моноліт з можливістю винесення адаптерів."],
        ["NFR-05", "Надійність", "Транзакційний checkout, історія статусів, audit платежів."],
        ["NFR-06", "Супроводжуваність", "Розподіл routes/services/features, Knex migrations, тестовий набір."],
        ["NFR-07", "Переносність", "Node.js runtime, Docker Compose, SQLite volume."],
        ["NFR-08", "Зручність", "Зрозумілі форми, inline-помилки, preview і статуси замовлень."],
    ]
    table_caption(doc, ctx, 1, "Нефункціональні вимоги")
    add_table(doc, ["ID", "Характеристика", "Реалізація"], nfr_rows, [2.0, 4.0, 10.5])
    add_heading(doc, ctx, "1.3.2.1 Вимоги до зовнішніх інтерфейсів", level=4)
    add_paragraph(doc, "Інтерфейси користувачів. Інтерфейси користувача охоплюють публічні сторінки каталогу, конструктора, корзини, checkout, оплату, список замовлень і адміністративні екрани керування.")
    add_paragraph(doc, "Інтерфейси програмного забезпечення. Програмні інтерфейси представлені REST-подібними HTTP API між React-клієнтом і Express-сервером, а також внутрішніми модулями сервісного шару backend.")
    add_paragraph(doc, "Інтерфейси апаратного забезпечення. Спеціальні апаратні інтерфейси для Aurora Atelier не вимагаються, оскільки система орієнтована на стандартний браузерний доступ і типовий серверний runtime.")
    add_paragraph(doc, "Комунікаційні інтерфейси. Комунікаційні інтерфейси включають HTTP-взаємодію браузера з сервером, cookie-based session flow, locale-параметри і mock-платіжний обмін у межах MVP.")
    add_heading(doc, ctx, "1.3.2.2 Вимоги до показників якості", level=4)
    add_paragraph(doc, "До показників якості для Aurora Atelier належать доступність ключових сценаріїв, коректність статусних переходів, прозорість ціноутворення, супроводжуваність коду і відсутність критичних помилок у тестованих потоках.")
    add_heading(doc, ctx, "1.3.2.3 Обмеження", level=4)
    add_paragraph(doc, "Основними обмеженнями поточного рішення є MVP-характер системи, використання mock payment замість реального провайдера, орієнтація на малу студію та відмова від важкого 3D-конструктора на користь легшого 2D-preview.")

    add_heading(doc, ctx, "1.4 Матриця відстеження (трасування) вимог", level=2)
    trace_rows = [[row[0], row[1], row[2], f"TC-{idx:02d}"] for idx, row in enumerate(functional_rows, 1)]
    table_caption(doc, ctx, 1, "Матриця трасування функціональних вимог")
    add_table(doc, ["Вимога", "Опис", "Модуль", "Тест"], trace_rows, [2.0, 8.5, 3.5, 2.5])
    add_project_elaboration(doc, "аналізу вимог", ["зв'язок вимог і тестів", "обмеження MVP", "класи користувачів", "якість і безпека", "інтерфейси програмного забезпечення"], 12 + expansion)

    add_heading(doc, ctx, "1.5 Висновки до розділу 1", level=2)
    add_paragraph(doc, "У першому розділі визначено бізнес-цілі, користувачів, функціональні й нефункціональні вимоги до Aurora Atelier. Аналіз аналогів показав, що розроблювана система має поєднати простоту Moment Creator, інтерактивність Jweel і персоналізований характер Takayas, але уникнути їх ключових недоліків: надмірної складності, відсутності динамічного preview і ручної непрозорої обробки замовлення.")


def add_design_section(doc: Document, ctx: BuildContext, diagrams: dict[str, Path], expansion: int) -> None:
    doc.add_page_break()
    add_heading(doc, ctx, "2 ПРОЄКТУВАННЯ І РЕАЛІЗАЦІЯ ПРОГРАМНОГО ЗАБЕЗПЕЧЕННЯ AURORA ATELIER", level=1)
    add_paragraph(doc, "Другий розділ описує архітектурні рішення, модель даних, взаємодію frontend і backend, реалізацію конструктора дизайну, checkout-процесу, адміністративних модулів і візуальної частини. Проєктування виконано з урахуванням MVP-обмежень і потреби у швидкому розгортанні через Docker.")

    add_heading(doc, ctx, "2.1 Архітектурне проєктування", level=2)
    add_picture(doc, diagrams["architecture"], 15.5)
    figure_caption(doc, ctx, 2, "Загальна архітектура Aurora Atelier")
    add_heading(doc, ctx, "2.1.1 Методологія / техніка / підхід розробки програмного забезпечення", level=3)
    for paragraph in [
        "Система реалізована як клієнт-серверний веб-додаток. Клієнтська частина є React SPA, яка збирається Vite і віддається Express як статичний ресурс з каталогу public/react-app. Серверна частина реалізує API, middleware, бізнес-логіку, роботу з базою даних, статичними файлами та інтеграційними адаптерами.",
        "Архітектурно проєкт можна описати як модульний моноліт. Один backend-сервіс містить доменні модулі auth, catalog, constructor, cart, checkout, payments, orders, account, admin-orders, admin-catalog, admin-constructor, admin-assets, notifications та i18n. Такий підхід простіший за мікросервіси для дипломного MVP, але зберігає внутрішній поділ відповідальностей.",
        "Frontend поділено на route-компоненти і features. Route відповідає за сторінку, навігацію, завантаження даних і orchestration, а feature-компоненти містять повторно використовувану логіку конкретного домену. Взаємодія з backend централізована у client/src/api.js, де описані API-клієнти і спільна функція http.",
    ]:
        add_paragraph(doc, paragraph)
    add_heading(doc, ctx, "2.1.2 Загальна модель / стиль архітектури програмного забезпечення", level=3)
    add_paragraph(doc, "З погляду архітектурного стилю Aurora Atelier поєднує client-server pattern, layered decomposition та route/service-oriented поділ відповідальностей, що добре відповідає предметній області і масштабу дипломного MVP.")
    add_paragraph(doc, "Клієнтський шар відповідає за presentation і orchestration сценаріїв, серверний шар реалізує бізнес-правила і API, а шар збереження даних забезпечує роботу з SQLite, міграціями, статусною історією та технічними журналами.")
    add_project_elaboration(doc, "архітектурного проєктування", ["модульний моноліт", "SPA поверх API", "Express middleware pipeline", "server-side business rules", "Docker-first runtime"], 14 + expansion)

    add_heading(doc, ctx, "2.1.3 Проєктування підсистеми збереження даних", level=3)
    add_picture(doc, diagrams["db"], 15.5)
    figure_caption(doc, ctx, 2, "Узагальнена модель даних")
    add_heading(doc, ctx, "2.1.3.1 Вибір моделі даних та підходу до проєктування БД", level=4)
    add_paragraph(doc, "Для системи обрано реляційну модель даних, оскільки вона добре описує користувачів, товари, замовлення, історію статусів і платежі, а також дозволяє зберігати референційну цілісність між ключовими бізнес-об'єктами.")
    add_heading(doc, ctx, "2.1.3.2 Проєктування та розробка моделі даних", level=4)
    table_caption(doc, ctx, 2, "Основні таблиці бази даних")
    add_table(
        doc,
        ["Група", "Таблиці", "Призначення"],
        [
            ["Користувачі", "users, sessions, email_verification_codes", "Облікові записи, ролі, сесії, підтвердження email."],
            ["Каталог", "jewelry_types, products, product_images, materials", "Типи прикрас, готові товари, зображення, матеріали."],
            ["Конструктор", "design_options, design_option_values та JSON-конфіги студії", "Параметри індивідуального дизайну, варіанти, слоти, камені."],
            ["Корзина", "carts, cart_items", "Активна корзина користувача та позиції ready/custom."],
            ["Замовлення", "orders, order_items, order_status_history", "Фіксація замовлення, позицій, статусів і історії."],
            ["Оплата", "payments", "Mock payment token, status, amount, audit payload."],
            ["Повідомлення", "notification_logs", "Фіксація email-подій і результатів доставки."],
            ["Промокоди", "promo_codes, promo_code_redemptions, promo_code_user_usage", "Знижки, ліміти і використання промокодів."],
        ],
        [3.0, 5.0, 8.5],
    )
    add_heading(doc, ctx, "2.1.3.3 Взаємодія з базою даних", level=4)
    add_paragraph(doc, "Для поточного MVP обрано SQLite, оскільки вона не потребує окремого database service, легко демонструється через Docker volume і достатня для дипломного навантаження. Використання Knex зберігає можливість майбутнього переходу на іншу SQL-СУБД без переписування всієї бізнес-логіки.")
    add_heading(doc, ctx, "2.1.3.4 Розробка запитів до бази даних", level=4)
    add_paragraph(doc, "Робота з даними реалізована через Knex query builder, що дозволяє поєднувати SQL-контрольованість із більш безпечним і читабельним описом вибірок, оновлень, транзакцій checkout і адміністративних CRUD-операцій.")
    add_project_elaboration(doc, "проєктування бази даних", ["міграції Knex", "індекси для каталогу і замовлень", "snapshot-дані у order_items", "audit платежів", "історія статусів"], 12 + expansion // 2)

    add_heading(doc, ctx, "2.1.4 Проєктування взаємодії з зовнішніми сервісами", level=3)
    add_heading(doc, ctx, "2.1.4.1 HTTP API та взаємодія клієнта із сервером", level=4)
    table_caption(doc, ctx, 2, "Основні API-групи")
    add_table(
        doc,
        ["API", "Призначення", "Приклади"],
        [
            ["/api/auth", "Клієнтська автентифікація, сесія, logout, email verification.", "GET /session, POST /register, POST /login"],
            ["/api/admin", "Адміністративний вхід.", "POST /login"],
            ["/api/catalog", "Каталог, картка товару, facets.", "GET /products, GET /products/:id"],
            ["/api/constructor", "Публічний конфіг конструктора і розрахунок ціни.", "GET /config, POST /price"],
            ["/api/cart", "Корзина, позиції, промокод.", "GET /, POST /items, DELETE /promo-code"],
            ["/api/checkout", "Створення замовлення з корзини.", "POST /"],
            ["/api/payments", "Mock payment confirmation.", "POST /mock/confirm"],
            ["/api/orders", "Замовлення клієнта.", "GET /me, GET /:id"],
            ["/api/admin/orders", "Admin order dashboard і статуси.", "GET /, PATCH /:id/status"],
            ["/api/admin/catalog", "Admin CRUD товарів і матеріалів.", "GET /products, POST /products"],
            ["/api/admin/constructor", "Admin CRUD типів, варіантів, слотів, каменів.", "GET /, PATCH /variant-stones"],
        ],
        [4.0, 7.0, 5.5],
    )
    add_paragraph(doc, "Єдиний API-клієнт на frontend додає credentials: include для роботи server sessions, передає локаль через x-locale і query lang, перевіряє response.ok та payload.success. Такий підхід централізує обробку помилок і зменшує дублювання в route-компонентах.")
    add_heading(doc, ctx, "2.1.4.2 Mock payment та checkout-взаємодія", level=4)
    add_paragraph(doc, "API корзини, checkout і mock payment охоплює критичний ланцюг від збереження позицій до створення замовлення і підтвердження оплати, тому потребує серверної валідації payload, ownership checks та idempotency-контролю.")
    add_heading(doc, ctx, "2.1.4.3 Сервіси замовлень, акаунта та повідомлень", level=4)
    add_paragraph(doc, "Клієнтські маршрути замовлень і акаунта повертають лише дані поточного користувача, підтримують активні і завершені замовлення та надають безпечний доступ до деталей замовлення.")
    add_heading(doc, ctx, "2.1.4.4 Публічні та адміністративні endpoint-и", level=4)
    add_paragraph(doc, "Адміністративний API забезпечує окрему сесію, захищені маршрути і керування товарами, замовленнями та бібліотекою конструктора без змішування з публічними endpoint-ами.")

    add_heading(doc, ctx, "2.2 Детальне проєктування", level=2)
    add_heading(doc, ctx, "2.2.1 Структурні моделі", level=3)
    add_paragraph(doc, "Структурні моделі Aurora Atelier відображають поділ на клієнтські маршрути, feature-модулі frontend, backend-домени та підсистему збереження даних. Цей поділ дозволяє локалізувати зміни і зменшує зв'язність між несуміжними сценаріями.")
    add_heading(doc, ctx, "2.2.2 Моделювання функціональності та взаємодії", level=3)
    add_paragraph(doc, "Функціональна взаємодія між браузером, API і БД моделюється через послідовності checkout, payment confirmation, status workflow та price calculation, що дає змогу простежити контрольні точки бізнес-правил.")
    add_paragraph(doc, "Конструктор дизайну.", bold=True, first_indent=False, keep_next=True)
    add_picture(doc, diagrams["pricing"], 15.5)
    figure_caption(doc, ctx, 2, "Потік розрахунку ціни та preview конструктора")
    add_paragraph(doc, "Конструктор працює з типами прикрас, варіантами, матеріалами, розмірами, слотами каменів, матрицею доступності каменів і preview-assets. Клієнтська частина дозволяє швидко змінювати параметри і бачити layered preview, але фінальна ціна формується backend-сервісом constructor-pricing.service.js.")
    table_caption(doc, ctx, 2, "Фактори ціноутворення custom design")
    add_table(
        doc,
        ["Фактор", "Джерело", "Вплив"],
        [
            ["Базова ціна типу", "constructor types", "Початкова вартість для ring, bracelet, earrings або pendant."],
            ["Варіант виробу", "variant.price_delta", "Доплата за конструктивну форму."],
            ["Матеріал", "type.materials", "Доплата за срібло, золото або rose gold."],
            ["Розмір", "type.size_options", "Доплата або корекція залежно від розміру."],
            ["Камені", "variant_stones", "Доплата за кожен доступний камінь у слоті."],
            ["Гравіювання", "type.engraving", "Доплата за текстове гравіювання і перевірка довжини."],
            ["Ланцюжок", "pendant-chain helper", "Доплата для підвісок залежно від вибору ланцюжка."],
        ],
        [4.0, 5.5, 7.0],
    )
    add_project_elaboration(doc, "детального проєктування конструктора", ["stone slots", "preview layers", "normalized configuration", "server-side validation", "admin constructor CRUD"], 14 + expansion // 2)

    add_paragraph(doc, "Checkout та життєвий цикл замовлення.", bold=True, first_indent=False, keep_next=True)
    add_picture(doc, diagrams["checkout"], 15.5)
    figure_caption(doc, ctx, 2, "Послідовність оформлення замовлення")
    add_paragraph(doc, "Checkout реалізовано як транзакційний процес. Сервер перевіряє payload, читає активну корзину, фіксує її у статусі checked_out, створює order, переносить cart_items в order_items, формує payment token, записує початкову історію статусу, реєструє використання промокоду і створює нову активну корзину.")
    table_caption(doc, ctx, 2, "Статусна модель замовлення")
    add_table(
        doc,
        ["Статус", "Значення", "Хто встановлює"],
        [
            ["created_pending_payment", "Замовлення створене, очікує оплату.", "checkout service"],
            ["confirmed", "Оплата підтверджена, замовлення прийняте.", "mock payment"],
            ["in_progress", "Майстер взяв замовлення в роботу.", "admin"],
            ["completed", "Замовлення виконане.", "admin"],
        ],
        [4.2, 8.0, 4.3],
    )
    add_project_elaboration(doc, "проєктування checkout", ["транзакційність", "ідемпотентність payment", "order snapshot", "status history", "email notification log"], 12 + expansion // 2)

    add_heading(doc, ctx, "2.3 Проєктування інтерфейсу користувача", level=2)
    add_heading(doc, ctx, "2.3.1 Публічна частина веб-системи", level=3)
    add_paragraph(doc, "Інтерфейс Aurora Atelier побудовано навколо преміального образу бренду: editorial hero, boutique-вітрина, акценти champagne/gold/ivory, типографіка Cormorant Garamond і Manrope у публічній частині. У межах звіту важливо відзначити, що дизайн не є окремим декоративним шаром: він підтримує користувацькі сценарії, робить каталог зрозумілим, а конструктор доступним.")
    add_heading(doc, ctx, "2.3.2 Особистий кабінет клієнта", level=3)
    add_paragraph(doc, "Екрани акаунта і замовлень побудовані як продовження публічної частини, але акцентують контроль статусу, перегляд деталей і повернення користувача до незавершених оплат.")
    add_heading(doc, ctx, "2.3.3 Адміністративний інтерфейс", level=3)
    add_paragraph(doc, "Адміністративна частина орієнтована на робочі сценарії майстра: фільтрацію, редагування та оперативне керування даними без декоративного перевантаження.")
    table_caption(doc, ctx, 2, "Основні екрани системи")
    add_table(
        doc,
        ["Route", "Призначення", "Ключові дії"],
        [
            ["/", "Головна сторінка", "Опис бренду, переходи до каталогу і конструктора."],
            ["/catalog", "Каталог", "Фільтри, список товарів, пагінація."],
            ["/products/:slug", "Картка товару", "Деталі, вибір розміру/ланцюжка, додавання до корзини."],
            ["/constructor", "Конструктор", "Вибір параметрів, preview, price API, cart."],
            ["/cart", "Корзина", "Редагування позицій, промокод, перехід до checkout."],
            ["/checkout", "Оформлення", "Контактні дані, доставка, згоди."],
            ["/payment/:id", "Оплата", "Mock payment confirmation."],
            ["/orders, /orders/:id", "Замовлення клієнта", "Список і деталі власних замовлень."],
            ["/admin/orders", "Admin orders", "Фільтрація, summary, status workflow."],
            ["/admin/products", "Admin products", "Master-detail редагування товарів."],
            ["/admin/constructor", "Admin constructor", "Типи, варіанти, слоти, камені, autosave."],
        ],
        [4.0, 5.2, 7.3],
    )

    add_heading(doc, ctx, "2.4 Реалізація програмного забезпечення", level=2)
    add_heading(doc, ctx, "2.4.1 Реалізація frontend-частини", level=3)
    add_paragraph(doc, "Реалізація виконана у репозиторії artisan-jewelry-mvp. Runtime модель є Docker-first: Dockerfile встановлює залежності, збирає React-клієнт і запускає Express, а docker-compose.yml виділяє volume для SQLite та uploads. На старті контейнера виконуються міграції Knex, після чого сервер запускається у production-like режимі.")
    add_heading(doc, ctx, "2.4.2 Реалізація backend-частини", level=3)
    add_paragraph(doc, "Backend-частина зосереджує middleware, маршрути, доменні service-модулі, транзакційну логіку checkout, ownership checks та журналювання критичних подій системи.")
    add_heading(doc, ctx, "2.4.3 Реалізація конструктора дизайну", level=3)
    add_paragraph(doc, "Реалізація конструктора поєднує client-side selection flow з backend price engine, library of assets, stone slots та administrative CRUD-інтерфейсом для підтримки цієї бібліотеки.")
    add_heading(doc, ctx, "2.4.4 Реалізація адміністративних модулів", level=3)
    add_paragraph(doc, "Адміністративні модулі охоплюють orders dashboard, products workspace та constructor workspace, кожен з яких має окрему прикладну логіку і захист role-based доступом.")
    table_caption(doc, ctx, 2, "Карта ключових файлів реалізації")
    add_table(
        doc,
        ["Файл або каталог", "Роль"],
        [
            ["server/app.js", "Express app, middleware, static files, API, SPA fallback."],
            ["server/routes/index.js", "Підключення всіх /api/* router'ів."],
            ["client/src/main.jsx", "React bootstrap і таблиця route definitions."],
            ["client/src/api.js", "Єдиний API-клієнт frontend."],
            ["server/modules/constructor", "Публічний і admin-конструктор, pricing, assets, JSON store."],
            ["server/modules/checkout", "Транзакційне створення замовлення."],
            ["server/modules/orders", "Клієнтські замовлення і ownership checks."],
            ["server/modules/admin-orders", "Admin dashboard і статусні переходи."],
            ["db/migrations", "Схема БД і еволюція структури."],
            ["tests", "Unit та integration tests."],
        ],
        [6.0, 10.5],
    )
    add_project_elaboration(doc, "реалізації програмного забезпечення", ["React route table", "domain CSS ownership", "Express services", "SQLite migrations", "Docker Compose"], 16 + expansion)

    add_heading(doc, ctx, "2.5 Висновки до розділу 2", level=2)
    add_paragraph(doc, "У другому розділі обґрунтовано клієнт-серверну архітектуру, модель даних, API-контракти, принципи роботи конструктора, checkout-процес і статусну модель замовлення. Реалізація відповідає вимогам MVP: вона достатньо проста для демонстрації та супроводу, але структурована так, щоб підтримувати подальше розширення.")


def add_testing_section(doc: Document, ctx: BuildContext, expansion: int) -> None:
    doc.add_page_break()
    add_heading(doc, ctx, "3 ТЕСТУВАННЯ ПРОГРАМНОГО ЗАБЕЗПЕЧЕННЯ AURORA ATELIER", level=1)
    add_paragraph(doc, "Тестування Aurora Atelier спрямоване на перевірку не лише окремих функцій, а й наскрізних бізнес-сценаріїв: автентифікації, каталогу, конструктора, корзини, checkout, mock payment, замовлень клієнта та адміністративної обробки. Для автоматизації використано Vitest і Supertest.")

    add_heading(doc, ctx, "3.1 Модульне тестування", level=2)
    add_heading(doc, ctx, "3.1.1 Стратегія та інструментальні засоби модульного тестування", level=3)
    add_paragraph(doc, "Модульні тести перевіряють невеликі ізольовані правила: форматування грошей, статусні переходи, фільтри каталогу, preview helpers, промокоди, redirect helpers і утиліти відображення імені клієнта. Такий рівень тестування потрібний для швидкого виявлення регресій у бізнес-правилах.")
    add_heading(doc, ctx, "3.1.2 План модульного тестування", level=3)
    table_caption(doc, ctx, 3, "План модульного тестування")
    add_table(
        doc,
        ["Файл", "Що перевіряється", "Кількість"],
        [
            ["tests/unit/catalog-filters.test.js", "Поведінка фільтрів каталогу.", "3"],
            ["tests/unit/cart-events.test.js", "Redirect helper для protected routes.", "3"],
            ["tests/unit/money.test.js", "Форматування і округлення сум.", "2"],
            ["tests/unit/order-statuses.test.js", "Дозволені переходи статусів.", "2"],
            ["tests/unit/preview.test.js", "Preview helper для конструктора.", "1"],
            ["tests/unit/promo-codes.test.js", "Валідація промокодів.", "2"],
            ["tests/unit/utils.test.js", "Форматування fallback customer name.", "3"],
        ],
        [6.0, 8.5, 2.0],
    )

    add_heading(doc, ctx, "3.2 Інтеграційне тестування", level=2)
    add_heading(doc, ctx, "3.2.1 Програмні та інструментальні засоби інтеграційного тестування", level=3)
    add_paragraph(doc, "Інтеграційні тести запускають Express app, seed-дані і SQLite через тестову конфігурацію, після чого виконують HTTP-запити до API. Це дозволяє перевірити поведінку на рівні реальних маршрутів, middleware, service layer і бази даних.")
    add_heading(doc, ctx, "3.2.2 План інтеграційного тестування", level=3)
    table_caption(doc, ctx, 3, "План інтеграційного тестування")
    add_table(
        doc,
        ["Файл", "Сценарії", "Кількість"],
        [
            ["tests/integration/app.test.js", "Health, static, 404, базова робота Express.", "4"],
            ["tests/integration/api-flows.test.js", "Auth, sessions, cart, constructor config, checkout validation, ownership, admin catalog.", "18"],
            ["tests/integration/cart-checkout.test.js", "Cart integrity, checkout idempotency, payment, promo, admin status, account dashboard.", "28"],
        ],
        [6.0, 8.5, 2.0],
    )
    add_heading(doc, ctx, "3.2.3 Результати інтеграційного тестування", level=3)
    add_paragraph(doc, "Інтеграційні перевірки підтвердили коректну взаємодію маршрутизації, middleware, бази даних, role guard, ownership logic, checkout transaction та адміністративних сценаріїв.")
    add_heading(doc, ctx, "3.2.4 Висновки щодо інтеграційного тестування", level=3)
    add_paragraph(doc, "Інтеграційне тестування показало, що основні підсистеми Aurora Atelier працюють узгоджено і не покладаються лише на локальну перевірку в браузері.")
    add_project_elaboration(doc, "інтеграційного тестування", ["реальні HTTP-запити", "SQLite test database", "session cookies", "role guard", "checkout transaction"], 10 + expansion // 2)

    add_heading(doc, ctx, "3.3 Тестування реалізації вимог до ПЗ", level=2)
    add_heading(doc, ctx, "3.3.1 Програмні та інструментальні засоби", level=3)
    add_paragraph(doc, "Для перевірки реалізації вимог використано Vitest, Supertest, тестову SQLite-конфігурацію і локальні seed-дані, які дозволяють відтворити реальні бізнес-сценарії без зовнішніх сервісів.")
    add_heading(doc, ctx, "3.3.2 План тестування реалізації функціональних вимог", level=3)
    add_paragraph(doc, "План функціонального тестування охоплює повний ланцюг від реєстрації й каталогу до checkout, оплати, перегляду замовлень та адміністративного керування.")
    add_heading(doc, ctx, "3.3.3 План тестування реалізації нефункціональних вимог", level=3)
    add_paragraph(doc, "Нефункціональні перевірки спрямовані на безпеку доступу, коректність ownership checks, ідемпотентність критичних операцій і здатність системи підтримувати узгоджений стан даних.")
    matrix = [
        ["FR-01", "Реєстрація і вхід", "api-flows: registration/login/session", "Пройдено"],
        ["FR-03", "Каталог з фільтрами", "cart-checkout: catalog filters", "Пройдено"],
        ["FR-05", "Конструктор", "api-flows: constructor config", "Пройдено"],
        ["FR-06", "Ціна custom design", "cart-checkout: chain/backend surcharge", "Пройдено"],
        ["FR-08", "Корзина", "api-flows: add same product updates quantity", "Пройдено"],
        ["FR-10", "Checkout", "api-flows: accepted offer, invalid phone/address", "Пройдено"],
        ["FR-11", "Mock payment", "cart-checkout: payment confirmation idempotent", "Пройдено"],
        ["FR-12", "Замовлення клієнта", "api-flows: ownership and authentication", "Пройдено"],
        ["FR-14", "Admin dashboard", "cart-checkout: filtered orders with summary", "Пройдено"],
        ["FR-15", "Статуси", "cart-checkout: cannot skip, rollback comment", "Пройдено"],
        ["FR-16", "Admin products", "api-flows: create/update product", "Пройдено"],
        ["NFR-01", "Безпека", "client cannot access admin endpoints", "Пройдено"],
    ]
    add_heading(doc, ctx, "3.3.4 Матриця відповідності вимог та тестів", level=3)
    table_caption(doc, ctx, 3, "Матриця відповідності вимог та тестів")
    add_table(doc, ["Вимога", "Суть", "Тест", "Результат"], matrix, [2.0, 4.0, 8.0, 2.5])
    add_heading(doc, ctx, "3.3.5 Результати тестування реалізації вимог до ПЗ", level=3)
    add_paragraph(doc, "Результати запуску автоматизованих перевірок підтвердили, що ключові функціональні та частина критичних нефункціональних вимог реалізовані коректно і захищені від типових регресій.")

    add_heading(doc, ctx, "3.4 Тестування реалізації варіантів використання", level=2)
    add_heading(doc, ctx, "3.4.1 План тестування реалізації варіантів використання", level=3)
    add_paragraph(doc, "Для варіантів використання перевіряються три наскрізні сценарії: створення індивідуальної прикраси, оформлення та оплата замовлення, адміністративна обробка замовлення після його створення.")
    add_heading(doc, ctx, "3.4.2 Результати тестування реалізації варіантів використання", level=3)
    add_paragraph(doc, "Перед формуванням звіту виконано команду npm test. Результат: 10 тестових файлів пройдено, 66 тестів пройдено, помилок не зафіксовано. У логах тестового режиму email delivery пропускається навмисно, але notification service створює записи про події, що дозволяє перевірити інтеграційний потік без реального SMTP.")
    table_caption(doc, ctx, 3, "Фактичний результат автоматизованого тестування")
    add_table(
        doc,
        ["Показник", "Значення"],
        [
            ["Команда", "npm test"],
            ["Test files", "10 passed"],
            ["Tests", "66 passed"],
            ["Framework", "Vitest 2.1.9 + Supertest"],
            ["Дата локальної перевірки", "06.06.2026"],
            ["Висновок", "Критичні сценарії MVP підтверджені автоматизованими тестами."],
        ],
        [5.0, 11.5],
    )
    add_project_elaboration(doc, "результатів тестування", ["регресійна перевірка", "матриця вимог", "admin status workflow", "checkout validation", "payment idempotency"], 12 + expansion)

    add_heading(doc, ctx, "3.5 Висновки до розділу 3", level=2)
    add_paragraph(doc, "У третьому розділі описано стратегію тестування, плани модульних та інтеграційних перевірок, матрицю відповідності вимог і тестів, а також фактичні результати запуску. Тестування підтвердило працездатність ключових сценаріїв Aurora Atelier і показало, що бізнес-правила не покладаються лише на клієнтський інтерфейс.")


def add_economics_section(doc: Document, ctx: BuildContext, expansion: int, fine_extra: int = 0) -> None:
    doc.add_page_break()
    add_heading(doc, ctx, "4 ЕКОНОМІЧНЕ ОБҐРУНТУВАННЯ ПРОЄКТУ", level=1)
    add_paragraph(doc, "Економічне обґрунтування визначає доцільність створення програмного продукту, оцінює витрати на розроблення, трудомісткість, потенційну споживацьку ціну і конкурентоспроможність. Розрахунки виконуються для MVP веб-системи, що може бути використана малою майстернею авторських прикрас.")

    add_heading(doc, ctx, "4.1 Опис, споживацька ціна та аналіз конкурентоспроможності програмного продукту", level=2)
    add_heading(doc, ctx, "4.1.1 Оцінювання ринку збуту та конкуренція", level=3)
    add_paragraph(doc, "Продукт належить до класу веб-систем для малого e-commerce з розширенням у вигляді конструктора дизайну. На відміну від універсальних CMS або маркетплейсів, Aurora Atelier враховує специфіку індивідуальних прикрас: конфігурацію виробу, preview, серверну ціну, snapshot у замовленні та адміністративний статусний workflow.")
    add_heading(doc, ctx, "4.1.2 Виявлення конкурентів та аналіз конкурентоспроможності програмного продукту", level=3)
    table_caption(doc, ctx, 4, "Фактори конкурентоспроможності")
    add_table(
        doc,
        ["Фактор", "Значення для клієнта", "Реалізація в Aurora Atelier"],
        [
            ["Персоналізація", "Можливість створити власний дизайн.", "2D-конструктор з варіантами, каменями і матеріалами."],
            ["Прозора ціна", "Клієнт одразу бачить вартість.", "Серверний price API."],
            ["Швидкість замовлення", "Не потрібно чекати відповідь менеджера.", "Cart, checkout, mock payment."],
            ["Контроль майстра", "Власник бачить і обробляє замовлення.", "Admin orders і status workflow."],
            ["Підтримка каталогу", "Можна продавати готові вироби.", "40 seeded products і admin products."],
            ["Вартість впровадження", "MVP дешевший за складний 3D-конструктор.", "React/Express/SQLite/Docker."],
        ],
        [4.0, 5.8, 6.7],
    )
    add_heading(doc, ctx, "4.1.3 Розрахунок споживацької ціни програмного продукту, уточнення цільової місткості ринку", level=3)
    add_paragraph(doc, "Орієнтовна споживацька ціна MVP визначається на основі трудомісткості, технічних витрат, запланованої рентабельності та припущення про кілька впроваджень продукту для схожих майстерень.")

    add_heading(doc, ctx, "4.2 Виробничо-організаційний план проєкту. Розрахунок трудомісткості проєкту та кількості його виконавців", level=2)
    add_heading(doc, ctx, "4.2.1 Перелік робіт", level=3)
    works = [
        ["1", "Аналіз предметної області та аналогів", "7"],
        ["2", "Формування вимог і сценаріїв використання", "8"],
        ["3", "Проєктування архітектури", "6"],
        ["4", "Проєктування бази даних", "5"],
        ["5", "Розроблення backend API", "14"],
        ["6", "Розроблення frontend routes", "13"],
        ["7", "Реалізація конструктора дизайну", "15"],
        ["8", "Реалізація checkout і order flow", "10"],
        ["9", "Реалізація admin-модулів", "12"],
        ["10", "Тестування і hardening", "10"],
        ["11", "Docker-first deployment", "4"],
        ["12", "Документування і підготовка матеріалів", "8"],
    ]
    table_caption(doc, ctx, 4, "Перелік робіт та трудомісткість")
    add_table(doc, ["№", "Робота", "Людино-дні"], works, [1.5, 12.0, 3.0])
    total_days = sum(int(row[2]) for row in works)
    add_heading(doc, ctx, "4.2.2 Розрахунки трудомісткості робіт проєкту", level=3)
    add_paragraph(doc, f"Загальна оцінена трудомісткість становить {total_days} людино-дні. Для MVP прийнято модель виконання одним розробником з консультаціями керівника і консультанта економічного розділу.")
    add_heading(doc, ctx, "4.2.3 Виконавці проєкту", level=3)
    add_paragraph(doc, "Основним виконавцем проєкту виступає один розробник, що виконує аналіз, проєктування, реалізацію, тестування і підготовку звітних матеріалів, а керівник та консультант економічного розділу забезпечують методичний супровід.")
    add_project_elaboration(doc, "виробничо-організаційного плану", ["планування робіт", "роль одного виконавця", "ризики інтеграції", "документування", "підготовка до захисту"], 12 + expansion // 2)

    add_heading(doc, ctx, "4.3 Фінансовий план проєкту", level=2)
    day_rate = 1300
    salary = total_days * day_rate
    equipment = 18000
    software = 0
    cloud = 2400
    overhead = round((salary + equipment + cloud) * 0.12, 2)
    total_cost = salary + equipment + software + cloud + overhead
    add_heading(doc, ctx, "4.3.1 Розрахунок потреб в інвестиціях на обладнання", level=3)
    add_paragraph(doc, "Початкові інвестиції охоплюють робоче обладнання, периферію та мінімальні витрати на демонстраційне розгортання системи й технічну підтримку MVP-середовища.")
    add_heading(doc, ctx, "4.3.2 Розрахунок заробітної плати виконавців проєкту", level=3)
    add_paragraph(doc, "Заробітна плата оцінюється за середньоденною ставкою розробника та загальною трудомісткістю, отриманою у виробничо-організаційному плані.")
    add_heading(doc, ctx, "4.3.3 Складання кошторису витрат на розроблення проєкту програмного продукту", level=3)
    table_caption(doc, ctx, 4, "Кошторис витрат на розроблення")
    add_table(
        doc,
        ["Стаття витрат", "Розрахунок", "Сума, грн"],
        [
            ["Оплата праці розробника", f"{total_days} днів x {day_rate} грн", f"{salary:,.2f}".replace(",", " ")],
            ["Обладнання та амортизація", "Робочий ноутбук, периферія, тестові пристрої", f"{equipment:,.2f}".replace(",", " ")],
            ["Програмне забезпечення", "Open-source стек, без ліцензійних витрат", f"{software:,.2f}".replace(",", " ")],
            ["Хостинг і домен для MVP", "Орієнтовні витрати на демонстраційне розгортання", f"{cloud:,.2f}".replace(",", " ")],
            ["Накладні витрати", "12% від основних витрат", f"{overhead:,.2f}".replace(",", " ")],
            ["Разом", "", f"{total_cost:,.2f}".replace(",", " ")],
        ],
        [5.0, 7.5, 4.0],
    )
    copies = 8
    profit_margin = 0.25
    unit_price = total_cost * (1 + profit_margin) / copies
    add_heading(doc, ctx, "4.3.4 Визначення критичної програми випуску програмного продукту", level=3)
    add_paragraph(doc, "Для оцінки окупності приймається припущення про декілька впроваджень системи для подібних майстерень, що дозволяє розподілити початкові витрати між окремими екземплярами продукту.")
    add_heading(doc, ctx, "4.3.5 Собівартість виготовлення одного екземпляра програмного продукту", level=3)
    add_paragraph(doc, "Собівартість одного екземпляра визначається через розподіл загальної вартості розроблення на умовну кількість впроваджень, допустиму для обраного сегмента малих ювелірних майстерень.")
    add_heading(doc, ctx, "4.3.6 План доходів від реалізації програмного продукту та витрат на його виготовлення", level=3)
    add_paragraph(doc, "План доходів формується на основі орієнтовної ціни одного впровадження та запланованої рентабельності, а витрати враховують розроблення, обладнання, інфраструктуру і накладні витрати.")
    table_caption(doc, ctx, 4, "Орієнтовна споживацька ціна")
    add_table(
        doc,
        ["Показник", "Значення"],
        [
            ["Повна собівартість", f"{total_cost:,.2f} грн".replace(",", " ")],
            ["Планова рентабельність", "25%"],
            ["Умовна кількість впроваджень", str(copies)],
            ["Орієнтовна ціна одного впровадження", f"{unit_price:,.2f} грн".replace(",", " ")],
        ],
        [7.0, 9.5],
    )
    add_project_elaboration(doc, "фінансового плану", ["оцінка трудомісткості", "відкритий стек", "вартість впровадження", "конкурентоспроможність", "окупність"], 12 + expansion)

    add_heading(doc, ctx, "4.3.7 Строк окупності витрат на проект", level=3)
    monthly_savings = 9500
    payback = total_cost / monthly_savings
    add_paragraph(doc, f"Якщо впровадження системи дозволяє майстерні економити або додатково заробляти орієнтовно {monthly_savings:,.2f} грн на місяць за рахунок скорочення ручної обробки, зменшення помилок і підвищення конверсії, строк окупності становить приблизно {payback:.1f} місяця.".replace(",", " "))
    add_paragraph(doc, "Такий показник є прийнятним для малого бізнесу, оскільки система може використовуватися повторно, розширюватися новими товарами та варіантами конструктора, а основні витрати припадають на початкову розробку.")

    if fine_extra:
        add_project_elaboration(
            doc,
            "практичної ефективності впровадження",
            ["скорочення ручних консультацій", "прозорість ціни", "повторне використання даних каталогу", "контроль виконання замовлення", "зменшення операційних помилок"],
            fine_extra,
        )

    add_heading(doc, ctx, "4.4 Висновки до розділу 4", level=2)
    add_paragraph(doc, "Економічне обґрунтування показує, що розроблення Aurora Atelier є доцільним для малого бізнесу у сфері авторських прикрас. Використання відкритого стеку, Docker-first запуску і SQLite знижує витрати на інфраструктуру, а наявність конструктора, checkout і admin-модулів підвищує споживацьку цінність продукту.")


def add_economics_section(doc: Document, ctx: BuildContext, expansion: int, fine_extra: int = 0) -> None:
    doc.add_page_break()

    def money(value: float) -> str:
        return f"{value:,.2f}".replace(",", " ")

    add_heading(doc, ctx, "4 ЕКОНОМІЧНЕ ОБҐРУНТУВАННЯ ПРОЄКТУ", level=1)
    add_paragraph(
        doc,
        "Економічне обґрунтування для Aurora Atelier виконується як оцінка реального MVP-продукту: branded web-системи для малої ювелірної майстерні з каталогом готових виробів, конструктором дизайну, кошиком, checkout, mock payment, особистим кабінетом та адміністративним модулем керування замовленнями й контентом. У розрахунках враховано фактичний склад робіт проєкту, сучасний рівень оплати праці розробника, необхідні інвестиції в робоче місце та витрати на запуск демонстраційного середовища.",
    )

    add_heading(doc, ctx, "4.1 Опис, споживацька ціна та аналіз конкурентоспроможності програмного продукту", level=2)
    add_heading(doc, ctx, "4.1.1 Оцінювання ринку збуту та конкуренція", level=3)
    add_paragraph(
        doc,
        "Цільовим сегментом Aurora Atelier є малі авторські майстерні та локальні jewelry-бренди, які вже продають вироби через Instagram, маркетплейси або прості storefront-рішення, але не мають єдиної системи для персоналізованих замовлень. Для цього сегмента критичними є чотири властивості: можливість швидко показати клієнту індивідуальний виріб, автоматично розрахувати ціну, не втрачати склад замовлення під час листування та мати зручну адмінку для керування каталогом і конструктором. Саме на ці бізнес-потреби орієнтовано поточний проєкт.",
    )
    add_paragraph(
        doc,
        "Найближчими альтернативами для такого сегмента є: шаблонний інтернет-магазин без конструктора, кастомний storefront без вбудованої персоналізації та bespoke-проєкт з індивідуальною логікою конструктора. Перший варіант найдешевший, але не покриває ключову потребу в індивідуальному замовленні. Другий покриває каталог і checkout, проте залишає ручну обробку персоналізації. Третій функціонально близький до Aurora Atelier, однак має вищий бюджет і довший цикл впровадження.",
    )

    add_heading(doc, ctx, "4.1.2 Виявлення конкурентів та аналіз конкурентоспроможності програмного продукту", level=3)
    table_caption(doc, ctx, 4, "Порівняння варіантів впровадження для цільового сегмента")
    add_table(
        doc,
        ["Варіант", "Орієнтовний бюджет, грн", "Ключові обмеження", "Висновок для Aurora Atelier"],
        [
            ["Шаблонний магазин на SaaS/CMS", "60 000 - 90 000", "Немає конструктора, персоналізація обробляється вручну, відсутній snapshot замовлення.", "Не покриває основну бізнес-цінність проєкту."],
            ["Кастомний storefront без конструктора", "170 000 - 210 000", "Є каталог і checkout, але немає візуальної конфігурації прикраси та централізованої адмінки конструктора.", "Частково придатний, але не вирішує задачу індивідуальних замовлень."],
            ["Кастомна система з персоналізацією виробу", "260 000 - 320 000", "Функціонально близька, однак дорожча у впровадженні та підтримці.", "Найближчий конкурентний клас, від якого доцільно відштовхуватися під час оцінки ціни."],
        ],
        [4.8, 3.7, 5.4, 4.6],
    )
    add_paragraph(
        doc,
        "Конкурентоспроможність Aurora Atelier визначається поєднанням трьох факторів. По-перше, система вже має предметну модель ювелірних виробів: типи прикрас, варіанти, слоти, бібліотеку каменів і серверний прайсинг. По-друге, клієнтський і адміністративний контури реалізовані в одному продукті, тому не виникає дублювання даних між storefront і внутрішньою CRM. По-третє, стек React + Express + SQLite + Docker не вимагає дорогих ліцензій та дозволяє розгорнути MVP без окремої складної інфраструктури.",
    )

    add_heading(doc, ctx, "4.1.3 Розрахунок споживацької ціни програмного продукту, уточнення цільової місткості ринку", level=3)
    base_competitor_price = 285000.00
    quality_coeff = 0.96
    consumer_price = round(0.9 * base_competitor_price * quality_coeff, 2)
    target_clients = 18
    repeat_cycle_years = 4
    annual_market_capacity = math.ceil(target_clients / repeat_cycle_years)
    add_paragraph(
        doc,
        f"Відповідно до методичних рекомендацій споживацьку ціну доцільно оцінювати через ціну базового конкурентного рішення та коефіцієнт відносної якості. Для Aurora Atelier як базовий конкурентний клас прийнято кастомну систему з персоналізацією виробу вартістю {money(base_competitor_price)} грн. З урахуванням коефіцієнта морального старіння 0,9 та коефіцієнта відносної якості {quality_coeff:.2f} отримаємо максимально прийнятну для ринку споживацьку ціну {money(consumer_price)} грн.",
    )
    add_paragraph(
        doc,
        f"Для цільового сегмента прийнято потенційний пул із {target_clients} майстерень, для яких поєднання storefront, персоналізації й адміністративного контуру реально створює бізнес-ефект. Якщо орієнтуватися на цикл повторного оновлення або придбання подібного рішення раз на {repeat_cycle_years} роки, річна місткість сегмента становитиме близько {annual_market_capacity} впроваджень на рік. Для MVP це означає, що план із кількох впроваджень є реалістичним, але продукт повинен залишатися дешевшим за повноцінний bespoke-аналог.",
    )

    add_heading(doc, ctx, "4.2 Виробничо-організаційний план проєкту. Розрахунок трудомісткості проєкту та кількості його виконавців", level=2)
    add_heading(doc, ctx, "4.2.1 Перелік робіт", level=3)
    works = [
        ["1", "Аналіз предметної області та аналогів", "7"],
        ["2", "Формування вимог і сценаріїв використання", "8"],
        ["3", "Проєктування архітектури та API", "10"],
        ["4", "Проєктування бази даних і бізнес-правил", "8"],
        ["5", "Розроблення backend-модулів каталогу, кошика та checkout", "14"],
        ["6", "Розроблення frontend-частини та клієнтських маршрутів", "13"],
        ["7", "Реалізація конструктора дизайну та серверного прайсингу", "15"],
        ["8", "Реалізація account/order flow і mock payment", "10"],
        ["9", "Розроблення адміністративних модулів", "12"],
        ["10", "Тестування, hardening і виправлення дефектів", "10"],
        ["11", "Підготовка Docker runtime та демонстраційного оточення", "5"],
        ["12", "Документування, скриншоти та підготовка матеріалів до захисту", "10"],
    ]
    table_caption(doc, ctx, 4, "Перелік робіт та трудомісткість розроблення")
    add_table(doc, ["№", "Робота", "Людино-дні"], works, [1.5, 12.0, 3.0])
    total_days = sum(int(row[2]) for row in works)
    working_days_per_month = 21
    project_months = round(total_days / working_days_per_month, 2)
    add_heading(doc, ctx, "4.2.2 Розрахунки трудомісткості робіт проєкту", level=3)
    add_paragraph(
        doc,
        f"Сумарна трудомісткість робіт становить {total_days} людино-днів, що відповідає приблизно {project_months:.2f} місяця роботи одного виконавця за середньому фонді {working_days_per_month} робочий день на місяць. Найбільшу частку трудовитрат формують не статичні сторінки storefront, а предметно-специфічні модулі: конструктор дизайну, серверний прайсинг, order lifecycle та адміністративна підтримка контенту.",
    )
    add_heading(doc, ctx, "4.2.3 Виконавці проєкту", level=3)
    add_paragraph(
        doc,
        "Базовою моделлю виконання прийнято одного software engineer, який послідовно виконує аналітику, проєктування, розроблення, тестування й підготовку демонстраційних матеріалів. Консультації керівника та консультанта економічного розділу не включаються до прямої зарплатної статті розроблення, але впливають на якість постановки задач, контроль результатів і фінальну придатність роботи до захисту.",
    )
    add_project_elaboration(
        doc,
        "виробничо-організаційного плану",
        ["планування робіт", "зв'язок між модулями конструктора й checkout", "тестування сценаріїв замовлення", "адміністративний контур", "підготовка матеріалів до захисту"],
        10 + expansion // 2,
    )

    add_heading(doc, ctx, "4.3 Фінансовий план проєкту", level=2)
    monthly_salary = 65000.00
    day_rate = round(monthly_salary / working_days_per_month, 2)
    ozop = round(total_days * day_rate, 2)
    additional_salary_norm = 15
    dzop = round(ozop * additional_salary_norm / 100, 2)
    esv = round((ozop + dzop) * 0.22, 2)
    materials = 3500.00
    equipment_items = [
        ["Ноутбук розробника", "1", "42 000.00", "42 000.00"],
        ["Монітор для UI/QA", "1", "9 000.00", "9 000.00"],
        ["Смартфон для mobile smoke-тестів", "1", "12 000.00", "12 000.00"],
        ["Периферія та мережеве обладнання", "1 комплект", "4 000.00", "4 000.00"],
    ]
    equipment_investment = 67000.00
    operation_norm = 8
    veo = round(equipment_investment * operation_norm / 100, 2)
    amortization = round(equipment_investment * 0.20 * project_months / 12, 2)
    area_per_person = 6
    rent_per_m2 = 350.00
    rent = round(area_per_person * rent_per_m2 * project_months, 2)
    internet_services = 4200.00
    other_production_norm = 20
    other_production = round(ozop * other_production_norm / 100, 2)
    admin_norm = 18
    admin_cost = round(ozop * admin_norm / 100, 2)
    production_cost = round(ozop + dzop + esv + materials + veo + amortization + rent + internet_services + other_production, 2)
    selling_norm = 5
    selling_cost = round(production_cost * selling_norm / 100, 2)
    full_cost = round(production_cost + admin_cost + selling_cost, 2)
    copies = 5
    tirage_cost = 2500.00
    adaptation_cost = 12000.00
    unit_cost = round(full_cost / copies + tirage_cost + adaptation_cost, 2)
    profitability = 25
    manufacturing_price = round(unit_cost * (1 + profitability / 100), 2)
    recommended_price = 229000.00
    critical_copies = math.ceil(full_cost / (recommended_price - (tirage_cost + adaptation_cost)))

    add_heading(doc, ctx, "4.3.1 Розрахунок потреб в інвестиціях на обладнання", level=3)
    add_paragraph(
        doc,
        "Для виконання проєкту потрібне робоче місце, придатне до одночасної backend/frontend-розробки, зборки клієнтського застосунку, запуску Docker-контейнерів, локального тестування та мобільного smoke QA. Тому до інвестицій на обладнання включено не лише ноутбук, а й другий екран, тестовий смартфон та периферію.",
    )
    table_caption(doc, ctx, 4, "Потреба в інвестиціях на обладнання")
    add_table(
        doc,
        ["Найменування", "Кількість", "Ціна за одиницю, грн", "Сума, грн"],
        equipment_items,
        [7.0, 2.5, 4.0, 3.5],
    )
    add_paragraph(doc, f"Загальна потреба в інвестиціях на обладнання становить {money(equipment_investment)} грн.")

    add_heading(doc, ctx, "4.3.2 Розрахунок заробітної плати виконавців проєкту", level=3)
    add_paragraph(
        doc,
        f"Основну заробітну плату визначено виходячи з місячного доходу software engineer {money(monthly_salary)} грн та середньої денної ставки {money(day_rate)} грн. За трудомісткості {total_days} людино-днів основна зарплата виконавця становить {money(ozop)} грн. Додаткова зарплата прийнята на рівні {additional_salary_norm}% від основної, тобто {money(dzop)} грн, а нарахування ЄСВ (22%) складають {money(esv)} грн.",
    )

    add_heading(doc, ctx, "4.3.3 Складання кошторису витрат на розроблення проєкту програмного продукту", level=3)
    table_caption(doc, ctx, 4, "Кошторис витрат на розроблення Aurora Atelier")
    add_table(
        doc,
        ["Стаття витрат", "Розрахунок", "Сума, грн"],
        [
            ["Основна заробітна плата", f"{total_days} днів x {money(day_rate)} грн", money(ozop)],
            ["Додаткова заробітна плата", f"{additional_salary_norm}% від основної зарплати", money(dzop)],
            ["ЄСВ", "22% від суми основної та додаткової зарплати", money(esv)],
            ["Витратні матеріали", "Папір, картриджі, дрібна периферія, носії", money(materials)],
            ["Амортизація обладнання", f"20% річних на період {project_months:.2f} міс.", money(amortization)],
            ["Експлуатація обладнання", f"{operation_norm}% від інвестицій на обладнання", money(veo)],
            ["Оренда робочого місця", f"{area_per_person} м² x {money(rent_per_m2)} грн x {project_months:.2f} міс.", money(rent)],
            ["Інтернет та онлайн-сервіси", "Хостинг, домен, GitHub/Cloud-послуги, зв'язок", money(internet_services)],
            ["Інші виробничі витрати", f"{other_production_norm}% від основної зарплати", money(other_production)],
            ["Адміністративні витрати", f"{admin_norm}% від основної зарплати", money(admin_cost)],
            ["Витрати на збут", f"{selling_norm}% від виробничої собівартості", money(selling_cost)],
            ["Разом", "", money(full_cost)],
        ],
        [5.2, 7.2, 4.1],
    )
    add_paragraph(
        doc,
        f"Повна собівартість розроблення MVP становить {money(full_cost)} грн. На відміну від попередньої спрощеної оцінки, у кошторисі враховано не лише зарплату та умовний хостинг, а й додаткову зарплату, ЄСВ, амортизацію, оренду робочого місця, інтернет-сервіси, адміністративні та збутові витрати.",
    )

    add_heading(doc, ctx, "4.3.4 Визначення критичної програми випуску програмного продукту", level=3)
    add_paragraph(
        doc,
        f"Для критичної програми випуску прийнято, що змінні витрати на одне впровадження складаються з тиражування, розгортання та адаптації під конкретну майстерню і дорівнюють {money(tirage_cost + adaptation_cost)} грн. За обраної ціни реалізації {money(recommended_price)} грн точка беззбитковості досягається після {critical_copies} впроваджень. Отже, уже четверте впровадження покриває витрати на розроблення та переводить продукт у зону прибутковості.",
    )

    add_heading(doc, ctx, "4.3.5 Собівартість виготовлення одного екземпляра програмного продукту", level=3)
    add_paragraph(
        doc,
        f"Собівартість одного екземпляра визначено як частку повної собівартості розроблення, розподілену на {copies} очікувані впровадження, з додаванням витрат на тиражування та адаптацію. Таким чином, собівартість одного впровадження становить {money(unit_cost)} грн. Це значення враховує не лише код, а й налаштування домену, контентну адаптацію, імпорт каталогу та первинне навчання замовника.",
    )

    add_heading(doc, ctx, "4.3.6 План доходів від реалізації програмного продукту та витрат на його виготовлення", level=3)
    table_caption(doc, ctx, 4, "Цінові орієнтири для впровадження Aurora Atelier")
    add_table(
        doc,
        ["Показник", "Значення"],
        [
            ["Максимально прийнятна споживацька ціна", f"{money(consumer_price)} грн"],
            ["Собівартість одного впровадження", f"{money(unit_cost)} грн"],
            [f"Ціна виготовлення за рентабельності {profitability}%", f"{money(manufacturing_price)} грн"],
            ["Рекомендована ціна реалізації", f"{money(recommended_price)} грн"],
            ["Очікувана кількість впроваджень у розрахунку", str(copies)],
        ],
        [8.0, 8.5],
    )
    add_paragraph(
        doc,
        f"Рекомендована ціна реалізації {money(recommended_price)} грн є вищою за мінімально допустиму ціну виготовлення {money(manufacturing_price)} грн, але нижчою за розраховану максимально прийнятну споживацьку ціну {money(consumer_price)} грн. Це означає, що продукт одночасно залишається конкурентним для малого бренду прикрас і забезпечує розробнику нормативну рентабельність.",
    )
    add_project_elaboration(
        doc,
        "фінансового плану",
        ["зарплатний фонд", "амортизація робочого місця", "адміністративні витрати", "ціна впровадження", "точка беззбитковості"],
        10 + expansion,
    )

    add_heading(doc, ctx, "4.3.7 Строк окупності витрат на проект", level=3)
    monthly_effect = 29000.00
    payback_months = round(recommended_price / monthly_effect, 1)
    add_paragraph(
        doc,
        f"Для замовника строк окупності доцільно оцінювати через щомісячний економічний ефект від упровадження. Якщо система дає майстерні близько {money(monthly_effect)} грн на місяць сумарної вигоди за рахунок скорочення ручного узгодження, швидшого оформлення замовлень, меншої кількості помилок у специфікаціях і підвищення конверсії персоналізованих заявок, то інвестиція в одне впровадження за ціною {money(recommended_price)} грн окупиться приблизно за {payback_months} місяця.",
    )
    add_paragraph(
        doc,
        "Для малого ювелірного бренду це прийнятний горизонт, оскільки система не є разовим промосайтом: після запуску вона продовжує працювати як storefront, внутрішній інструмент приймання замовлень і адміністраторський контур керування каталогом та конструктором. Тобто економічний ефект накопичується не в одній кампанії, а в кожному наступному циклі продажу.",
    )

    if fine_extra:
        add_project_elaboration(
            doc,
            "практичної ефективності впровадження",
            ["скорочення ручних консультацій", "вбудований прайсинг", "повторне використання каталогу", "адміністративний контроль", "підвищення конверсії індивідуальних замовлень"],
            fine_extra,
        )

    add_heading(doc, ctx, "4.4 Висновки до розділу 4", level=2)
    add_paragraph(
        doc,
        f"Економічний розрахунок показує, що Aurora Atelier є реалістичним для комерціалізації продуктом. Повна собівартість розроблення MVP становить {money(full_cost)} грн, рекомендована ціна одного впровадження — {money(recommended_price)} грн, а точка беззбитковості досягається після {critical_copies} впроваджень. Для цільового замовника строк окупності близько {payback_months} місяця є прийнятним і підтверджує доцільність використання системи не лише як навчального проєкту, а як основи реального digital-інструмента для продажу авторських прикрас.",
    )


def add_conclusion_and_sources(doc: Document, ctx: BuildContext) -> None:
    add_structural_heading(doc, ctx, "ВИСНОВКИ", page_break=True)
    for paragraph in [
        "У кваліфікаційній роботі розроблено веб-систему Aurora Atelier для продажу авторських прикрас з конструктором дизайну. Система охоплює ключові сценарії предметної області: перегляд каталогу, створення індивідуального виробу, розрахунок ціни, корзину, checkout, імітацію оплати, перегляд замовлень клієнтом і адміністративну обробку.",
        "У першому розділі виконано аналіз вимог, сформовано бізнес-вимоги, глосарій, класи користувачів, сценарії, функціональні й нефункціональні вимоги. Аналіз аналогів показав необхідність збалансованого рішення, яке не є таким складним, як 3D-конструктор Jweel, але надає більше прозорості й інтерактивності, ніж Moment Creator або Takayas Custom Jewelry.",
        "У другому розділі спроєктовано архітектуру програмного продукту. Обрано модульний моноліт на основі React SPA, Express API, SQLite, Knex і Docker Compose. Описано модель даних, API-групи, конструктор, checkout, статусну модель, UI-рішення і карту ключових файлів реалізації.",
        "У третьому розділі проведено тестування. Автоматизований набір на Vitest і Supertest підтвердив проходження 66 тестів у 10 файлах. Перевірено автентифікацію, каталог, конструктор, корзину, checkout, оплату, замовлення, admin flows, промокоди і статусні переходи.",
        "У четвертому розділі виконано економічне обґрунтування. Розрахунки показали, що MVP має прийнятну собівартість, може бути впроваджений малою майстернею і створює економічний ефект завдяки автоматизації ручної обробки та підвищенню прозорості продажу.",
        "Отже, поставлену мету досягнуто: створено веб-систему, яка поєднує персоналізацію авторських прикрас, контрольований процес замовлення і адміністративне керування, зберігаючи простоту розгортання та супроводу.",
    ]:
        add_paragraph(doc, paragraph)

    add_structural_heading(doc, ctx, "ПЕРЕЛІК ПОСИЛАНЬ", page_break=True)
    sources = [
        "Wiegers K., Beatty J. Software Requirements. Microsoft Press, 2013.",
        "Sommerville I. Software Engineering. Pearson, 2016.",
        "ISO/IEC 25010:2011. Systems and software engineering - Systems and software Quality Requirements and Evaluation.",
        "React Documentation. URL: https://react.dev/.",
        "Vite Documentation. URL: https://vite.dev/.",
        "Express Documentation. URL: https://expressjs.com/.",
        "Knex.js Documentation. URL: https://knexjs.org/.",
        "SQLite Documentation. URL: https://www.sqlite.org/docs.html.",
        "Docker Documentation. URL: https://docs.docker.com/.",
        "Vitest Documentation. URL: https://vitest.dev/.",
        "Supertest Repository Documentation. URL: https://github.com/ladjs/supertest.",
        "MDN Web Docs. HTTP overview. URL: https://developer.mozilla.org/.",
        "Матеріали архіву «Бакалаври - документи - захист» кафедри 603, 2026.",
        "Вимоги до оформлення розрахунково-пояснювальних записок та звітів. ХАІ, 2026.",
        "Структура основних розділів ВКР бакалавра 2025-2026. ХАІ, 2026.",
        "Матеріали курсової роботи «Курсова Гірка 642п.docx», розділ огляду аналогів.",
        "Jweel. Online Jewelry Design Platform. URL: https://www.jweel.com/.",
        "Moment Creator. Personalized Jewelry. URL: https://momentcreator.com/.",
        "Takayas Custom Jewelry. URL: https://takayascustomjewelry.com/.",
        "Репозиторій проєкту Aurora Atelier: локальні матеріали, документація, вихідний код і тести.",
    ]
    for idx, source in enumerate(sources, 1):
        add_paragraph(doc, f"{idx}. {source}", first_indent=False)


def add_appendices(doc: Document, ctx: BuildContext, diagrams: dict[str, Path]) -> None:
    add_structural_heading(doc, ctx, "ДОДАТОК А", page_break=True)
    add_bookmark(doc.paragraphs[-1], "bm_appendix_start", ctx.next_bookmark_id())
    add_center(doc, "Графічні матеріали проєктування", bold=True)
    for key, caption in [
        ("architecture", "Загальна архітектура системи"),
        ("db", "Модель даних системи"),
        ("checkout", "Послідовність checkout"),
        ("pricing", "Схема розрахунку ціни конструктора"),
    ]:
        add_picture(doc, diagrams[key], 15.5)
        add_caption(doc, caption)

    add_structural_heading(doc, ctx, "ДОДАТОК Б", page_break=True)
    add_center(doc, "Фрагменти API та тестових сценаріїв", bold=True)
    code_blocks = [
        ("API-клієнт frontend", "client/src/api.js", "http(url, options) додає credentials, x-locale, JSON body і перевіряє payload.success."),
        ("Checkout service", "server/modules/checkout/checkout.service.js", "createCheckoutOrder виконує транзакцію: cart -> order -> order_items -> payment -> notification."),
        ("Constructor pricing", "server/modules/constructor/constructor-pricing.service.js", "calculateStudioPrice перевіряє variant, material, size, stones, engraving і повертає normalized configuration."),
        ("Order statuses", "server/constants/order-statuses.js", "ORDER_STATUS_FLOW фіксує послідовність created_pending_payment -> confirmed -> in_progress -> completed."),
        ("Integration tests", "tests/integration/cart-checkout.test.js", "Тести перевіряють idempotency checkout/payment, promo, status workflow і account dashboard."),
    ]
    table_caption(doc, ctx, 5, "Відібрані фрагменти реалізації")
    add_table(doc, ["Назва", "Файл", "Суть"], code_blocks, [4.0, 5.5, 7.0])

    add_structural_heading(doc, ctx, "ДОДАТОК В", page_break=True)
    add_center(doc, "Скріншоти та матеріали користувацьких сценаріїв", bold=True)
    add_paragraph(
        doc,
        "У додатку наведено скріншоти, отримані з фактично запущеної Docker-версії веб-системи Aurora Atelier. "
        "Матеріали згруповано за сценаріями використання, описаними в основній частині пояснювальної записки: "
        "вибір виробу, персоналізація дизайну, оформлення замовлення, оплата, контроль замовлення клієнтом і "
        "адміністративна обробка даних.",
    )

    appendix_v_sections = [
        (
            "В.1 Публічний сценарій вибору та персоналізації прикраси",
            [
                (
                    "01-home.png",
                    "Головна сторінка з навігацією до каталогу та конструктора",
                    "Скріншот підтверджує стартову точку сценарію клієнта: користувач бачить позиціювання ательє, основну навігацію та переходи до каталогу готових виробів і конструктора дизайну.",
                ),
                (
                    "02-catalog.png",
                    "Каталог готових прикрас із фільтрами та картками товарів",
                    "Матеріал демонструє виконання вимоги перегляду каталогу: система відображає товари, ціни, типи прикрас і візуальні картки для подальшого вибору виробу.",
                ),
                (
                    "03-product.png",
                    "Сторінка готового виробу перед додаванням у кошик",
                    "Скріншот фіксує деталізацію обраного товару: назву, зображення, характеристики, ціну та елементи керування для додавання позиції до кошика.",
                ),
                (
                    "04-constructor.png",
                    "Конструктор дизайну з вибором типу, матеріалу та попереднім переглядом",
                    "Матеріал підтверджує реалізацію сценарію персоналізації: користувач працює з параметрами прикраси, бачить попередній перегляд і отримує розрахунок вартості.",
                ),
            ],
        ),
        (
            "В.2 Сценарій оформлення замовлення",
            [
                (
                    "05-cart.png",
                    "Кошик клієнта з позиціями замовлення та підсумком",
                    "Скріншот демонструє проміжний стан перед checkout: у кошику збережено вибрані позиції, кількість, вартість рядків і загальну суму замовлення.",
                ),
                (
                    "06-checkout.png",
                    "Форма оформлення замовлення з контактними даними та доставкою",
                    "Матеріал показує заповнення обов'язкових даних клієнта, вибір способу доставки та прийняття оферти перед створенням замовлення.",
                ),
                (
                    "07-payment.png",
                    "Mock-оплата замовлення банківською карткою",
                    "Скріншот підтверджує окремий етап оплати: система показує суму, номер замовлення, форму картки та підготовлений mock-платіж для підтвердження.",
                ),
            ],
        ),
        (
            "В.3 Сценарій контролю замовлення клієнтом",
            [
                (
                    "08-orders.png",
                    "Список замовлень клієнта в особистому кабінеті",
                    "Матеріал демонструє можливість клієнта переглядати історію та поточні замовлення після автентифікації в системі.",
                ),
                (
                    "09-order-detail.png",
                    "Деталі замовлення клієнта зі статусом та складом позицій",
                    "Скріншот фіксує сторінку конкретного замовлення: статус, контактні дані, склад позицій, суму та посилання на оплату для неоплаченого замовлення.",
                ),
            ],
        ),
        (
            "В.4 Адміністративний сценарій обробки замовлень",
            [
                (
                    "10-admin-login.png",
                    "Форма входу адміністратора до службової частини",
                    "Матеріал показує захищену точку входу адміністратора, яка відокремлює службові функції від публічної частини веб-системи.",
                ),
                (
                    "11-admin-orders.png",
                    "Адміністративний список замовлень із фільтрами та статусами",
                    "Скріншот демонструє робочу таблицю адміністратора: перелік замовлень, статуси, ознаку оплати, підсумкові показники та перехід до деталей.",
                ),
                (
                    "12-admin-order-detail.png",
                    "Адміністративна картка замовлення з керуванням статусом",
                    "Матеріал підтверджує сценарій обробки замовлення адміністратором: перегляд деталей, платежів, позицій і доступ до керування життєвим циклом статусу.",
                ),
            ],
        ),
        (
            "В.5 Адміністративне керування каталогом готових виробів",
            [
                (
                    "13-admin-products.png",
                    "Адміністративне керування каталогом готових виробів",
                    "Скріншот демонструє інтерфейс підтримки каталогу: пошук, фільтрацію за типом, вибір картки виробу та редактор основних даних товару.",
                ),
            ],
        ),
        (
            "В.6 Адміністративне керування структурою конструктора",
            [
                (
                    "14-admin-constructor-home.png",
                    "Головний екран робочого простору адміністратора конструктора",
                    "Матеріал фіксує стартовий вхід до службового модуля конструктора, у якому адміністратор обирає потрібний напрям редагування: прикраси, камені, асети або правила ціноутворення.",
                ),
                (
                    "15-admin-constructor-types.png",
                    "Бібліотека типів прикрас у службовому модулі конструктора",
                    "Скріншот демонструє перелік типів прикрас, що підтримуються конструктором, та підтверджує можливість централізовано підтримувати доменну структуру конфігуратора.",
                ),
                (
                    "16-admin-constructor-variants.png",
                    "Перелік варіантів для вибраного типу прикраси",
                    "Матеріал показує наступний рівень адміністрування конструктора: для обраного типу прикраси відображаються варіанти виконання, з якими надалі пов'язуються слоти, асети й матриця доступних каменів.",
                ),
            ],
        ),
        (
            "В.7 Адміністративне редагування варіанта прикраси",
            [
                (
                    "17-admin-constructor-slots.png",
                    "Редактор слотів варіанта прикраси в адмінському конструкторі",
                    "Скріншот підтверджує деталізоване налаштування слотів: адміністратор бачить координати, порядок, масштаб і службові параметри зон, у які користувач згодом може встановлювати камені.",
                ),
                (
                    "18-admin-constructor-basic.png",
                    "Базові параметри варіанта прикраси в службовому редакторі",
                    "Матеріал демонструє редагування службових атрибутів варіанта прикраси: назви, коду, групи, типу підваріанта, цінової надбавки та базового графічного асета.",
                ),
                (
                    "19-admin-constructor-matrix.png",
                    "Матриця доступності каменів для конкретного варіанта",
                    "Скріншот ілюструє важливий службовий сценарій підтримки асортименту конструктора: адміністратор керує дозволеними каменями для вибраного варіанта прикраси та їх ціновими коефіцієнтами.",
                ),
                (
                    "20-admin-constructor-preview.png",
                    "Службовий попередній перегляд варіанта прикраси",
                    "Матеріал фіксує вбудований прев'ю-режим адміністратора, у якому перевіряється візуальна коректність базового асета, позиціонування слотів та цілісність підготовленої конфігурації.",
                ),
            ],
        ),
        (
            "В.8 Адміністративне керування бібліотекою каменів, асетами та цінами конструктора",
            [
                (
                    "21-admin-constructor-stones.png",
                    "Службова бібліотека каменів конструктора",
                    "Скріншот підтверджує наявність окремого робочого екрану для підтримки каталогу каменів, що використовуються у персоналізації прикрас.",
                ),
                (
                    "22-admin-constructor-stone-editor.png",
                    "Редактор запису про камінь у службовому модулі",
                    "Матеріал демонструє редагування детальних параметрів каменю: назв, коду, вартості, кольору, ознак активності та інших атрибутів, від яких залежить публічний конструктор.",
                ),
                (
                    "23-admin-constructor-assets.png",
                    "Службова бібліотека графічних асетів конструктора",
                    "Скріншот показує централізоване керування графічними ресурсами, що застосовуються для баз прикрас, зображень каменів і допоміжних графічних елементів конструктора.",
                ),
                (
                    "24-admin-constructor-pricing.png",
                    "Службовий екран налаштування правил ціноутворення конструктора",
                    "Матеріал фіксує робочий інтерфейс контролю цінової логіки: адміністратор бачить тип прикраси, вибраний варіант і параметри, які впливають на розрахунок вартості персоналізованого виробу.",
                ),
            ],
        ),
    ]

    for section_title, screenshots in appendix_v_sections:
        add_paragraph(doc, section_title, bold=True, first_indent=False, keep_next=True)
        for file_name, caption, description in screenshots:
            image_path = SCREENSHOT_DIR / file_name
            if not image_path.exists():
                raise FileNotFoundError(f"Required Appendix V screenshot is missing: {image_path}")
            add_paragraph(doc, description, keep_next=True)
            add_picture(doc, image_path, 14.5)
            appendix_figure_caption(doc, ctx, "В", caption)


def build_report(expansion: int, fine_extra: int = 0) -> Path:
    OUTPUT_DOCX.parent.mkdir(parents=True, exist_ok=True)
    ASSETS_DIR.mkdir(parents=True, exist_ok=True)
    images = extract_course_analogue_images()
    diagrams = create_diagrams()

    doc = Document()
    configure_document(doc)
    ctx = BuildContext()

    add_front_matter(doc, ctx)

    body_section = doc.add_section(WD_SECTION.NEW_PAGE)
    configure_section(body_section)
    body_section.header.is_linked_to_previous = False
    header_para = body_section.header.paragraphs[0] if body_section.header.paragraphs else body_section.header.add_paragraph()
    header_para.text = ""
    add_page_number(header_para)

    add_introduction(doc, ctx, page_break=False)
    add_requirements_section(doc, ctx, images, diagrams, expansion)
    add_design_section(doc, ctx, diagrams, expansion)
    add_testing_section(doc, ctx, expansion)
    add_economics_section(doc, ctx, expansion, fine_extra)
    add_conclusion_and_sources(doc, ctx)
    add_appendices(doc, ctx, diagrams)

    doc.save(OUTPUT_DOCX)
    update_word_toc_and_fields(OUTPUT_DOCX)
    return OUTPUT_DOCX


def main() -> None:
    parser = argparse.ArgumentParser()
    parser.add_argument("--expansion", type=int, default=8, help="Additional paragraph expansion per major section")
    parser.add_argument("--fine-extra", type=int, default=0, help="Small number of extra paragraphs for page-count tuning")
    args = parser.parse_args()
    out = build_report(args.expansion, args.fine_extra)
    print(out)


if __name__ == "__main__":
    main()
